"""LLM-ready model documentation package builder."""

from __future__ import annotations

import csv
import json
from collections.abc import Mapping
from datetime import UTC, datetime
from io import BytesIO, StringIO
from pathlib import Path
from time import perf_counter
from typing import Any
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

import pandas as pd

from quant_pd_framework.decision_summary import build_decision_summary
from quant_pd_framework.figure_exports import FigureExportAsset

PACKAGE_ROOT = "llm_documentation_package"
PACKAGE_VERSION = "1.6"
START_HERE_DIR = f"{PACKAGE_ROOT}/00_START_HERE"
DOCUMENT_TEMPLATE_DIR = f"{PACKAGE_ROOT}/01_document_template"
DOCUMENT_PLANNING_DIR = f"{PACKAGE_ROOT}/02_document_planning"
DOCX_WORKFLOW_DIR = f"{PACKAGE_ROOT}/03_docx_workflow"
EVIDENCE_DIR = f"{PACKAGE_ROOT}/04_evidence"
VISUAL_ASSETS_DIR = f"{PACKAGE_ROOT}/05_visual_assets"
VALIDATION_CONTROLS_DIR = f"{PACKAGE_ROOT}/06_validation_controls"
OPERATOR_INSTRUCTIONS_DIR = f"{PACKAGE_ROOT}/07_operator_prompts"
TOOLS_DIR = f"{PACKAGE_ROOT}/08_tools"
GENERATED_OUTPUTS_DIR = f"{PACKAGE_ROOT}/09_generated_outputs"
SOURCE_ARTIFACTS_DIR = f"{EVIDENCE_DIR}/source_artifacts"
CONVERTED_TABLES_DIR = f"{EVIDENCE_DIR}/converted_tables"
TABLE_PREVIEWS_DIR = f"{EVIDENCE_DIR}/table_previews"
FIGURE_ASSETS_ROOT = f"{VISUAL_ASSETS_DIR}/figures"
DEFAULT_MAX_INCLUDED_FILE_BYTES = 25 * 1024 * 1024
DEFAULT_MAX_TABLE_PREVIEW_ROWS = 250
DEFAULT_MAX_TABLE_PREVIEW_COLUMNS = 80
DEFAULT_MAX_CHART_FILES = 60

ROW_LEVEL_ARTIFACT_KEYS = {
    "input_snapshot",
    "input_snapshot_csv",
    "input_snapshot_parquet",
    "predictions",
    "predictions_csv",
    "predictions_parquet",
    "full_data_predictions",
    "large_data_training_sample",
}
MODEL_BINARY_ARTIFACT_KEYS = {"model"}
HEAVY_DIRECTORY_ARTIFACT_KEYS = {
    "code_snapshot_dir",
    "monitoring_bundle_dir",
    "sample_development_dir",
    "full_data_scoring_dir",
    "large_data_metadata_dir",
}
HEAVY_REPORT_NAMES = {"interactive_report.html"}
LLM_TEXT_SUFFIXES = {".csv", ".json", ".md", ".txt", ".html"}
LLM_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".svg"}
DOCUMENT_FIGURE_PRIORITY_PATTERNS = (
    ("roc", "auc"),
    ("ks",),
    ("calibration",),
    ("threshold",),
    ("lift", "gain"),
    ("feature_importance", "driver", "importance"),
    ("partial_dependence", "pdp"),
    ("scorecard", "woe", "bin"),
    ("actual_vs_predicted", "residual"),
    ("psi", "stability"),
    ("backtest",),
    ("quantile",),
)


def build_llm_documentation_package(
    snapshot: Mapping[str, Any],
    *,
    max_included_file_bytes: int = DEFAULT_MAX_INCLUDED_FILE_BYTES,
    max_chart_files: int = DEFAULT_MAX_CHART_FILES,
) -> bytes:
    """Builds a zipped evidence package for downstream LLM document drafting."""

    payload = build_llm_documentation_context_payload(snapshot)
    return build_llm_documentation_package_from_payload(
        payload,
        max_included_file_bytes=max_included_file_bytes,
        max_chart_files=max_chart_files,
    )


def build_llm_documentation_context_payload(snapshot: Mapping[str, Any]) -> dict[str, Any]:
    """Returns a JSON-safe context payload used by the package builder and UI cache."""

    diagnostics_tables = _mapping(snapshot.get("diagnostics_tables"))
    summary = build_decision_summary(snapshot)
    table_previews = {
        str(name): _frame_payload(table)
        for name, table in diagnostics_tables.items()
        if isinstance(table, pd.DataFrame) and not _looks_row_level_table(str(name), table)
    }
    return {
        "package_version": PACKAGE_VERSION,
        "created_at_utc": datetime.now(UTC).isoformat(),
        "run": {
            "run_id": snapshot.get("run_id", ""),
            "execution_mode": snapshot.get("execution_mode", ""),
            "model_type": snapshot.get("model_type", ""),
            "target_mode": snapshot.get("target_mode", ""),
            "target_column": snapshot.get("target_column", ""),
            "labels_available": bool(snapshot.get("labels_available", False)),
            "input_shape": dict(_mapping(snapshot.get("input_shape"))),
            "feature_summary": dict(_mapping(snapshot.get("feature_summary"))),
            "split_summary": _json_safe(snapshot.get("split_summary", {})),
            "run_timing": _json_safe(snapshot.get("run_timing", {})),
            "run_diagnostics": _json_safe(snapshot.get("run_diagnostics", {})),
        },
        "decision_summary": {
            "recommendation": summary["recommendation"],
            "level": summary["level"],
            "rationale": list(summary["rationale"]),
            "cards": _records_payload(pd.DataFrame(summary["cards"])),
            "metrics": _records_payload(summary["metric_frame"]),
            "issues": _records_payload(summary["issue_frame"]),
            "feature_drivers": _records_payload(summary["feature_frame"]),
            "validation_checklist": _records_payload(summary["validation_checklist_frame"]),
            "traceability_map": _records_payload(summary["traceability_frame"]),
        },
        "metrics": _json_safe(snapshot.get("metrics", {})),
        "statistical_tests": _json_safe(snapshot.get("statistical_tests", {})),
        "warnings": [str(warning) for warning in snapshot.get("warnings", [])],
        "feature_columns": [str(column) for column in snapshot.get("feature_columns", [])],
        "numeric_features": [str(column) for column in snapshot.get("numeric_features", [])],
        "categorical_features": [
            str(column) for column in snapshot.get("categorical_features", [])
        ],
        "feature_importance": _frame_payload(snapshot.get("feature_importance")),
        "diagnostic_table_previews": table_previews,
        "diagnostic_table_inventory": [
            {
                "table_name": str(name),
                "rows": int(table.shape[0]),
                "columns": int(table.shape[1]),
                "included_in_context": str(name) in table_previews,
            }
            for name, table in diagnostics_tables.items()
            if isinstance(table, pd.DataFrame)
        ],
        "artifacts": {
            str(key): str(value)
            for key, value in _mapping(snapshot.get("artifacts")).items()
            if value
        },
    }


def build_llm_documentation_package_from_payload(
    payload: Mapping[str, Any],
    *,
    max_included_file_bytes: int = DEFAULT_MAX_INCLUDED_FILE_BYTES,
    max_chart_files: int = DEFAULT_MAX_CHART_FILES,
    generated_chart_assets: tuple[FigureExportAsset, ...] = (),
    build_profile: list[dict[str, Any]] | None = None,
) -> bytes:
    """Builds a zipped LLM documentation package from a JSON-safe payload."""

    package_started = perf_counter()
    profile_rows = list(build_profile or [])
    output = BytesIO()
    included: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []
    added_arcnames: set[str] = set()
    artifacts = _mapping(payload.get("artifacts"))
    run_root = _resolve_run_root(artifacts)

    with ZipFile(output, mode="w", compression=ZIP_DEFLATED, compresslevel=3) as archive:
        _write_text(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/README_LLM_PACKAGE.md",
            _build_readme(),
            included,
            evidence_area="orientation",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/START_HERE.md",
            _build_start_here(),
            included,
            evidence_area="orientation",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{START_HERE_DIR}/PACKAGE_MAP.md",
            _build_package_map(),
            included,
            evidence_area="orientation",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/context/model_document_context.json",
            payload,
            included,
            evidence_area="llm_context",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/context/model_document_context.md",
            _build_context_markdown(payload),
            included,
            evidence_area="llm_context",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{START_HERE_DIR}/MODEL_FACTS_DIGEST.md",
            _build_model_facts_digest(payload),
            included,
            evidence_area="model_facts_digest",
            source="generated",
            llm_use=(
                "High-confidence run facts for LLM drafting. Cite underlying evidence "
                "files when making final document claims."
            ),
        )
        _write_json(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/regulatory/sr_11_7_evidence_mapping.json",
            _build_guidance_mapping(),
            included,
            evidence_area="regulatory_mapping",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{OPERATOR_INSTRUCTIONS_DIR}/prompt_generate_model_methodology.md",
            _build_prompt_template(),
            included,
            evidence_area="llm_prompt",
            source="generated",
            llm_use=(
                "Operator prompt for controlled LLM drafting. Do not cite as model evidence."
            ),
        )
        _write_text(
            archive,
            added_arcnames,
            f"{OPERATOR_INSTRUCTIONS_DIR}/THREE_PROMPTS_FOR_LLM_USE.txt",
            _build_three_prompt_text(),
            included,
            evidence_area="llm_prompt",
            source="generated",
            llm_use=(
                "Copy/paste prompt sequence for controlled LLM documentation drafting. "
                "Do not cite as model evidence."
            ),
        )
        _write_text(
            archive,
            added_arcnames,
            f"{OPERATOR_INSTRUCTIONS_DIR}/prompts/prompt_create_docx_methodology.md",
            _build_docx_prompt_template(),
            included,
            evidence_area="llm_prompt",
            source="generated",
            llm_use=(
                "DOCX-specific operator prompt for controlled LLM document generation. "
                "Do not cite as model evidence."
            ),
        )
        _write_text(
            archive,
            added_arcnames,
            f"{DOCUMENT_PLANNING_DIR}/default_model_methodology_outline.md",
            _build_default_outline(),
            included,
            evidence_area="document_outline",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/DOCX_BUILD_INSTRUCTIONS.md",
            _build_docx_build_instructions(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use="DOCX authoring instructions for LLM or code-enabled document generation.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/MODEL_DOCUMENT_STYLE_GUIDE.md",
            _build_model_document_style_guide(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use="Formatting and style guide for model methodology DOCX output.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/DOCX_QUALITY_CHECKLIST.md",
            _build_docx_quality_checklist(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use="Final DOCX quality checklist for human review and LLM self-checking.",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/resolved_toc_binding.schema.json",
            _build_resolved_toc_binding_schema(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/resolved_toc_binding_instructions.md",
            _build_resolved_toc_binding_instructions(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use="Instructions for mapping a custom TOC to Quant Studio evidence.",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/methodology_doc_spec.schema.json",
            _build_methodology_doc_spec_schema(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/methodology_doc_spec_example.json",
            _build_methodology_doc_spec_example(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{GENERATED_OUTPUTS_DIR}/README_GENERATED_OUTPUTS.md",
            _build_generated_outputs_readme(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use="Output workspace instructions for generated TOC bindings, specs, and DOCX.",
        )
        section_map = _build_section_evidence_map(payload)
        _write_json(
            archive,
            added_arcnames,
            f"{DOCUMENT_PLANNING_DIR}/document_section_evidence_map.json",
            section_map,
            included,
            evidence_area="section_evidence_map",
            source="generated",
        )
        _write_csv(
            archive,
            added_arcnames,
            f"{DOCUMENT_PLANNING_DIR}/document_section_evidence_map.csv",
            section_map,
            included,
            evidence_area="section_evidence_map",
            source="generated",
            llm_use="Maps methodology document sections to exact package evidence sources.",
        )
        _write_section_evidence_folders(
            archive=archive,
            added_arcnames=added_arcnames,
            section_map=section_map,
            included=included,
        )
        regulatory_crosswalk = _build_regulatory_crosswalk(payload)
        _write_json(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/regulatory/regulatory_documentation_crosswalk.json",
            regulatory_crosswalk,
            included,
            evidence_area="regulatory_crosswalk",
            source="generated",
        )
        _write_csv(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/regulatory/regulatory_documentation_crosswalk.csv",
            regulatory_crosswalk,
            included,
            evidence_area="regulatory_crosswalk",
            source="generated",
            llm_use="Detailed SR 11-7/OCC-style crosswalk for model-document drafting.",
        )
        approved_claims = _build_approved_claims(payload)
        _write_json(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/claims/approved_claims.json",
            approved_claims,
            included,
            evidence_area="approved_claims",
            source="generated",
        )
        _write_csv(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/claims/approved_claims.csv",
            approved_claims,
            included,
            evidence_area="approved_claims",
            source="generated",
            llm_use="Evidence-backed claim library to reduce hallucinated document language.",
        )
        target_document_schema = _build_target_document_schema(payload, section_map)
        _write_json(
            archive,
            added_arcnames,
            f"{DOCUMENT_PLANNING_DIR}/target_document_schema.json",
            target_document_schema,
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/evidence_strength_policy.json",
            _build_evidence_strength_policy(),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/document_completion_rules.json",
            _build_document_completion_rules(),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/controlled_vocabulary.json",
            _build_controlled_vocabulary(),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/draft_validation_rules.json",
            _build_draft_validation_rules(),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{DOCUMENT_PLANNING_DIR}/template_binding.json",
            _build_template_binding(section_map),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/llm_redaction_policy.json",
            _build_llm_redaction_policy(),
            included,
            evidence_area="document_control",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/document_quality_rubric.json",
            _build_document_quality_rubric(),
            included,
            evidence_area="document_quality",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/document_quality_rubric.md",
            _build_document_quality_rubric_markdown(),
            included,
            evidence_area="document_quality",
            source="generated",
            llm_use="Reviewer rubric for scoring LLM-generated methodology drafts.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/citation_coverage_validator.md",
            _build_citation_coverage_validator(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Citation coverage checks to apply to LLM-generated drafts.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/unsupported_claim_detector.md",
            _build_unsupported_claim_detector(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Unsupported-claim screening rules for LLM-generated drafts.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/regulatory_language_guardrails.md",
            _build_regulatory_language_guardrails(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Language restrictions for regulatory model-risk documentation.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{TOOLS_DIR}/validate_llm_draft.py",
            _build_llm_draft_validator_script(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Standalone no-dependency draft citation and unsupported-claim checker.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{TOOLS_DIR}/validate_docx_spec.py",
            _build_docx_spec_validator_script(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Validates the structured methodology document spec before DOCX build.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{TOOLS_DIR}/build_docx_from_spec.py",
            _build_docx_builder_script(),
            included,
            evidence_area="docx_authoring_control",
            source="generated",
            llm_use=(
                "Builds model_methodology.docx deterministically from "
                "methodology_doc_spec.json."
            ),
        )
        _write_text(
            archive,
            added_arcnames,
            f"{TOOLS_DIR}/validate_docx_output.py",
            _build_docx_output_validator_script(),
            included,
            evidence_area="draft_validation",
            source="generated",
            llm_use="Checks that the generated DOCX exists and contains expected document parts.",
        )
        documentation_gaps = _build_documentation_gaps(payload)
        _write_csv(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/gaps/documentation_gaps.csv",
            documentation_gaps,
            included,
            evidence_area="documentation_gaps",
            source="generated",
            llm_use="Missing evidence register that tells the LLM what not to invent.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/gaps/documentation_gaps.md",
            _build_documentation_gaps_markdown(documentation_gaps),
            included,
            evidence_area="documentation_gaps",
            source="generated",
            llm_use="Plain-English missing evidence register for model-document drafting.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/methodology/model_type_writing_guide.md",
            _build_model_type_writing_guide(payload),
            included,
            evidence_area="model_type_writing_guide",
            source="generated",
            llm_use="Model-family-specific methodology writing guidance.",
        )
        _write_prompt_variants(archive, added_arcnames, included)
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/citation_rules.md",
            _build_citation_rules(),
            included,
            evidence_area="citation_rules",
            source="generated",
            llm_use="Strict citation rules for every factual LLM-generated claim.",
        )
        _write_tone_profiles(archive, added_arcnames, included)
        _write_text(
            archive,
            added_arcnames,
            f"{VALIDATION_CONTROLS_DIR}/human_review_checklist.md",
            _build_human_review_checklist(),
            included,
            evidence_area="human_review_checklist",
            source="generated",
            llm_use="Human verification checklist for any LLM-generated document.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/interpretation_briefs/feature_dictionary_narrative.md",
            _build_feature_dictionary_narrative(payload),
            included,
            evidence_area="feature_dictionary_narrative",
            source="generated",
            llm_use="Prose-ready summary of selected features, lineage, and documentation gaps.",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/interpretation_briefs/metrics_interpretation_brief.md",
            _build_metrics_interpretation_brief(payload),
            included,
            evidence_area="metrics_interpretation_brief",
            source="generated",
            llm_use="Prose-ready interpretation of captured model metrics.",
        )
        _write_template_files(archive, added_arcnames, included)
        _write_diagnostic_table_previews(archive, added_arcnames, payload, included)
        _include_run_artifacts(
            archive=archive,
            added_arcnames=added_arcnames,
            artifacts=artifacts,
            run_root=run_root,
            included=included,
            skipped=skipped,
            max_included_file_bytes=max_included_file_bytes,
            max_chart_files=max_chart_files,
        )
        _include_generated_chart_assets(
            archive=archive,
            added_arcnames=added_arcnames,
            generated_chart_assets=generated_chart_assets,
            included=included,
            skipped=skipped,
            max_included_file_bytes=max_included_file_bytes,
        )
        figure_placement_rows = _build_figure_placement_manifest(included)
        _write_csv(
            archive,
            added_arcnames,
            f"{VISUAL_ASSETS_DIR}/figure_placement_manifest.csv",
            figure_placement_rows,
            included,
            evidence_area="figure_placement_manifest",
            source="generated",
            llm_use=(
                "DOCX figure placement plan. Use chart file paths as visual evidence, "
                "but cite the underlying chart/table evidence named in this manifest."
            ),
        )
        _write_json(
            archive,
            added_arcnames,
            f"{VISUAL_ASSETS_DIR}/figure_placement_manifest.json",
            figure_placement_rows,
            included,
            evidence_area="figure_placement_manifest",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{VISUAL_ASSETS_DIR}/FIGURE_CONTACT_SHEET.md",
            _build_figure_contact_sheet(figure_placement_rows),
            included,
            evidence_area="figure_placement_manifest",
            source="generated",
            llm_use="Human and LLM-readable visual index for document figure selection.",
        )
        table_placement_rows = _build_table_placement_manifest(included)
        _write_csv(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/table_placement_manifest.csv",
            table_placement_rows,
            included,
            evidence_area="table_placement_manifest",
            source="generated",
            llm_use="DOCX table placement plan for controlled methodology drafting.",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{DOCX_WORKFLOW_DIR}/table_placement_manifest.json",
            table_placement_rows,
            included,
            evidence_area="table_placement_manifest",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{EVIDENCE_DIR}/interpretation_briefs/chart_interpretation_brief.md",
            _build_chart_interpretation_brief(included, skipped),
            included,
            evidence_area="chart_interpretation_brief",
            source="generated",
            llm_use="Chart inventory and interpretation guardrails for document drafting.",
        )
        profile_rows.append(
            {
                "stage": "assemble_llm_package",
                "elapsed_seconds": round(perf_counter() - package_started, 3),
                "included_file_count_before_profile": len(included),
                "skipped_file_count_before_profile": len(skipped),
                "note": (
                    "Package assembly timing excludes the final bytes returned to the "
                    "Streamlit download widget."
                ),
            }
        )
        _write_json(
            archive,
            added_arcnames,
            f"{START_HERE_DIR}/llm_package_build_profile.json",
            profile_rows,
            included,
            evidence_area="package_build_profile",
            source="generated",
        )
        checklist_rows = _build_evidence_checklist(included, skipped)
        _write_csv(
            archive,
            added_arcnames,
            f"{START_HERE_DIR}/llm_evidence_checklist.csv",
            checklist_rows,
            included,
            evidence_area="evidence_checklist",
            source="generated",
        )
        _write_text(
            archive,
            added_arcnames,
            f"{START_HERE_DIR}/llm_evidence_checklist.md",
            _build_checklist_markdown(checklist_rows),
            included,
            evidence_area="evidence_checklist",
            source="generated",
        )
        evidence_index_rows = _build_evidence_index(included, skipped)
        _write_csv(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/EVIDENCE_INDEX.csv",
            evidence_index_rows,
            included,
            evidence_area="evidence_index",
            source="generated",
            llm_use=(
                "Authoritative file classification index. Only rows with "
                "cite_as_evidence=true should support factual document claims."
            ),
        )
        _write_text(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/DO_NOT_CITE.md",
            _build_do_not_cite_markdown(evidence_index_rows),
            included,
            evidence_area="citation_boundary",
            source="generated",
            llm_use=(
                "Explicit list of package files that guide workflow but must not "
                "be cited as model evidence."
            ),
        )
        _write_csv(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/source_citation_map.csv",
            _citable_records(included),
            included,
            evidence_area="citation_map",
            source="generated",
        )
        _write_json(
            archive,
            added_arcnames,
            f"{PACKAGE_ROOT}/llm_evidence_manifest.json",
            {
                "package_version": PACKAGE_VERSION,
                "created_at_utc": datetime.now(UTC).isoformat(),
                "run_id": _mapping(payload.get("run")).get("run_id", ""),
                "usage": (
                    "Evidence package for LLM-assisted model methodology drafting. "
                    "LLM output must be reviewed by qualified model-development and "
                    "validation personnel."
                ),
                "privacy_policy": (
                    "Row-level input data, row-level predictions, serialized model binaries, "
                    "and full code snapshots are excluded by default."
                ),
                "included_files": included,
                "skipped_files": skipped,
                "file_classification_policy": (
                    "Use EVIDENCE_INDEX.csv as the authoritative citation boundary. "
                    "Only files marked cite_as_evidence=true should support factual "
                    "model-document claims."
                ),
            },
            included,
            evidence_area="manifest",
            source="generated",
        )

    return output.getvalue()


def _include_run_artifacts(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    artifacts: Mapping[str, Any],
    run_root: Path | None,
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
    max_chart_files: int,
) -> None:
    for key, value in artifacts.items():
        path = Path(str(value))
        if not path.exists():
            skipped.append(
                _skip_record(
                    key=key,
                    original_path=path,
                    reason="artifact path was not found on disk",
                )
            )
            continue
        if key in ROW_LEVEL_ARTIFACT_KEYS:
            skipped.append(
                _skip_record(
                    key=key,
                    original_path=path,
                    reason="row-level input or prediction data is excluded",
                )
            )
            continue
        if key in MODEL_BINARY_ARTIFACT_KEYS:
            skipped.append(
                _skip_record(
                    key=key,
                    original_path=path,
                    reason="serialized model binaries are excluded",
                )
            )
            continue
        if key in HEAVY_DIRECTORY_ARTIFACT_KEYS:
            skipped.append(
                _skip_record(
                    key=key,
                    original_path=path,
                    reason="large supporting directory is excluded",
                )
            )
            continue
        if path.is_dir():
            if key == "tables_dir":
                _include_table_directory(
                    archive=archive,
                    added_arcnames=added_arcnames,
                    path=path,
                    run_root=run_root,
                    included=included,
                    skipped=skipped,
                    max_included_file_bytes=max_included_file_bytes,
                )
            elif key == "figures_dir":
                _include_figure_directory(
                    archive=archive,
                    added_arcnames=added_arcnames,
                    path=path,
                    run_root=run_root,
                    included=included,
                    skipped=skipped,
                    max_included_file_bytes=max_included_file_bytes,
                    max_chart_files=max_chart_files,
                )
            else:
                skipped.append(
                    _skip_record(
                        key=key,
                        original_path=path,
                        reason="directory is not part of the LLM evidence allowlist",
                    )
                )
            continue
        _include_file(
            archive=archive,
            added_arcnames=added_arcnames,
            key=key,
            path=path,
            run_root=run_root,
            included=included,
            skipped=skipped,
            max_included_file_bytes=max_included_file_bytes,
        )


def _include_generated_chart_assets(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    generated_chart_assets: tuple[FigureExportAsset, ...],
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
) -> None:
    for asset in generated_chart_assets:
        if not asset.arcname.startswith(f"{PACKAGE_ROOT}/"):
            skipped.append(
                _skip_record(
                    key="generated_figures",
                    original_path=Path(asset.arcname),
                    reason="generated chart asset was outside the LLM package namespace",
                    size_bytes=asset.size_bytes,
                )
            )
            continue
        if asset.arcname in added_arcnames:
            continue
        if asset.size_bytes > max_included_file_bytes:
            skipped.append(
                _skip_record(
                    key="generated_figures",
                    original_path=Path(asset.arcname),
                    reason=(
                        "generated chart asset exceeds package file-size cap "
                        f"({max_included_file_bytes} bytes)"
                    ),
                    size_bytes=asset.size_bytes,
                )
            )
            continue
        archive.writestr(
            asset.arcname,
            asset.data,
            compress_type=_compress_type_for_generated_asset(asset),
        )
        added_arcnames.add(asset.arcname)
        included.append(
            {
                "package_path": asset.arcname,
                "original_path": "generated_from_completed_run_visualizations",
                "artifact_key": "generated_figures",
                "evidence_area": "visual_evidence",
                "source": "generated_chart_export",
                "size_bytes": asset.size_bytes,
                "llm_use": _generated_chart_llm_use(asset),
            }
        )


def _generated_chart_llm_use(asset: FigureExportAsset) -> str:
    if asset.file_format == "png":
        return (
            "Chart image generated from the completed run for selective insertion "
            "into LLM-drafted model documentation."
        )
    if asset.file_format == "html":
        return "Interactive chart file generated from the completed run for visual review."
    if asset.file_format == "json":
        return "Manifest listing generated individual chart files and export warnings."
    return "Supporting file for generated chart evidence."


def _compress_type_for_generated_asset(asset: FigureExportAsset) -> int:
    if asset.file_format == "png":
        return ZIP_STORED
    return ZIP_DEFLATED


def _include_table_directory(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    path: Path,
    run_root: Path | None,
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
) -> None:
    for table_path in sorted(path.rglob("*")):
        if not table_path.is_file():
            continue
        if _looks_row_level_path(table_path):
            skipped.append(
                _skip_record(
                    key="tables_dir",
                    original_path=table_path,
                    reason="row-level table path is excluded",
                )
            )
            continue
        if table_path.suffix.lower() == ".parquet":
            _include_parquet_table_as_csv(
                archive=archive,
                added_arcnames=added_arcnames,
                path=table_path,
                run_root=run_root,
                included=included,
                skipped=skipped,
                max_included_file_bytes=max_included_file_bytes,
            )
            continue
        if table_path.suffix.lower() in {".csv", ".json", ".md", ".txt"}:
            _include_file(
                archive=archive,
                added_arcnames=added_arcnames,
                key="tables_dir",
                path=table_path,
                run_root=run_root,
                included=included,
                skipped=skipped,
                max_included_file_bytes=max_included_file_bytes,
            )


def _include_figure_directory(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    path: Path,
    run_root: Path | None,
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
    max_chart_files: int,
) -> None:
    chart_count = 0
    for figure_path in sorted(path.rglob("*")):
        if not figure_path.is_file():
            continue
        suffix = figure_path.suffix.lower()
        if suffix not in LLM_IMAGE_SUFFIXES and suffix != ".html":
            continue
        if chart_count >= max_chart_files:
            skipped.append(
                _skip_record(
                    key="figures_dir",
                    original_path=figure_path,
                    reason=f"chart file cap of {max_chart_files} was reached",
                )
            )
            continue
        if _file_size(figure_path) > max_included_file_bytes:
            skipped.append(
                _skip_record(
                    key="figures_dir",
                    original_path=figure_path,
                    reason=(
                        "chart file exceeds package file-size cap "
                        f"({max_included_file_bytes} bytes)"
                    ),
                )
            )
            continue
        chart_count += 1
        _include_file(
            archive=archive,
            added_arcnames=added_arcnames,
            key="figures_dir",
            path=figure_path,
            run_root=run_root,
            included=included,
            skipped=skipped,
            max_included_file_bytes=max_included_file_bytes,
        )


def _include_file(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    key: str,
    path: Path,
    run_root: Path | None,
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
) -> None:
    file_name = path.name
    suffix = path.suffix.lower()
    if file_name in HEAVY_REPORT_NAMES:
        skipped.append(
            _skip_record(
                key=key,
                original_path=path,
                reason="large standalone interactive report is referenced but not duplicated",
            )
        )
        return
    if suffix not in LLM_TEXT_SUFFIXES and suffix not in LLM_IMAGE_SUFFIXES and suffix != ".py":
        skipped.append(
            _skip_record(
                key=key,
                original_path=path,
                reason=f"file suffix `{suffix or 'none'}` is not LLM package friendly",
            )
        )
        return
    size = _file_size(path)
    if size > max_included_file_bytes:
        skipped.append(
            _skip_record(
                key=key,
                original_path=path,
                reason=f"file exceeds package file-size cap ({max_included_file_bytes} bytes)",
                size_bytes=size,
            )
        )
        return
    arcname = _source_arcname(path, run_root)
    if arcname in added_arcnames:
        return
    archive.write(path, arcname)
    added_arcnames.add(arcname)
    included.append(
        {
            "package_path": arcname,
            "original_path": str(path),
            "artifact_key": key,
            "evidence_area": _evidence_area_for_path(path, key),
            "source": "run_artifact",
            "size_bytes": size,
            "llm_use": _llm_use_for_path(path, key),
        }
    )


def _include_parquet_table_as_csv(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    path: Path,
    run_root: Path | None,
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    max_included_file_bytes: int,
) -> None:
    if _file_size(path) > max_included_file_bytes:
        skipped.append(
            _skip_record(
                key="tables_dir",
                original_path=path,
                reason="parquet table exceeds conversion file-size cap",
            )
        )
        return
    try:
        frame, original_rows, original_columns = _read_parquet_table_preview(path)
    except Exception as exc:  # pragma: no cover - defensive against parquet engine issues.
        skipped.append(
            _skip_record(
                key="tables_dir",
                original_path=path,
                reason=f"parquet table could not be converted to CSV: {exc}",
            )
        )
        return
    if _looks_row_level_table(path.stem, frame, total_rows=original_rows):
        skipped.append(
            _skip_record(
                key="tables_dir",
                original_path=path,
                reason="row-level parquet table is excluded",
            )
        )
        return
    csv_text = frame.to_csv(index=False)
    relative = _relative_source_path(path, run_root).with_suffix(".csv")
    arcname = f"{CONVERTED_TABLES_DIR}/{_posix(relative)}"
    if arcname in added_arcnames:
        return
    archive.writestr(arcname, csv_text)
    added_arcnames.add(arcname)
    included.append(
        {
            "package_path": arcname,
            "original_path": str(path),
            "artifact_key": "tables_dir",
            "evidence_area": "diagnostic_tables",
            "source": "converted_parquet_table",
            "size_bytes": len(csv_text.encode("utf-8")),
            "llm_use": (
                "Converted diagnostic table preview for LLM-readable evidence review. "
                f"Included {frame.shape[0]:,} of {original_rows:,} rows and "
                f"{frame.shape[1]:,} of {original_columns:,} columns."
            ),
        }
    )


def _read_parquet_table_preview(path: Path) -> tuple[pd.DataFrame, int, int]:
    """Reads only the LLM preview window from a Parquet table when possible."""

    try:
        import pyarrow.parquet as pq
    except ImportError:  # pragma: no cover - pyarrow is expected with parquet support.
        frame = pd.read_parquet(path)
        original_rows, original_columns = frame.shape
        return (
            frame.iloc[
                :DEFAULT_MAX_TABLE_PREVIEW_ROWS,
                :DEFAULT_MAX_TABLE_PREVIEW_COLUMNS,
            ].copy(),
            int(original_rows),
            int(original_columns),
        )

    parquet_file = pq.ParquetFile(path)
    metadata = parquet_file.metadata
    original_rows = int(metadata.num_rows) if metadata is not None else 0
    column_names = list(parquet_file.schema_arrow.names)
    selected_columns = column_names[:DEFAULT_MAX_TABLE_PREVIEW_COLUMNS]
    frames: list[pd.DataFrame] = []
    remaining_rows = DEFAULT_MAX_TABLE_PREVIEW_ROWS
    for row_group_index in range(parquet_file.num_row_groups):
        if remaining_rows <= 0:
            break
        table = parquet_file.read_row_group(row_group_index, columns=selected_columns)
        if table.num_rows == 0:
            continue
        row_group_frame = table.to_pandas()
        frames.append(row_group_frame.head(remaining_rows))
        remaining_rows -= min(len(row_group_frame), remaining_rows)

    if frames:
        preview = pd.concat(frames, ignore_index=True)
    else:
        preview = pd.DataFrame(columns=selected_columns)
    return preview, original_rows, len(column_names)


def _write_diagnostic_table_previews(
    archive: ZipFile,
    added_arcnames: set[str],
    payload: Mapping[str, Any],
    included: list[dict[str, Any]],
) -> None:
    table_previews = _mapping(payload.get("diagnostic_table_previews"))
    for table_name, table_payload in table_previews.items():
        rows = list(_mapping(table_payload).get("records", []))
        if not rows:
            continue
        safe_name = _safe_name(str(table_name))
        _write_csv(
            archive,
            added_arcnames,
            f"{TABLE_PREVIEWS_DIR}/{safe_name}.csv",
            rows,
            included,
            evidence_area="diagnostic_table_preview",
            source="generated",
            llm_use=(
                "LLM-readable preview of a diagnostic table captured in the completed run."
            ),
        )


def _write_template_files(
    archive: ZipFile,
    added_arcnames: set[str],
    included: list[dict[str, Any]],
) -> None:
    template_files = {
        f"{DOCUMENT_TEMPLATE_DIR}/DROP_TABLE_OF_CONTENTS_HERE.md": (
            "# Drop Your Table Of Contents Here\n\n"
            "Replace this file with your institution's model technical document table "
            "of contents before sending the package to an LLM. The prompt instructs the "
            "LLM to fill that structure using only package evidence.\n"
        ),
        f"{DOCUMENT_TEMPLATE_DIR}/table_of_contents_instructions.md": (
            "# Table Of Contents Instructions\n\n"
            "Use this folder for a group-specific model document structure. If a custom "
            "table of contents is present, instruct the LLM to preserve the section order, "
            "headings, and terminology while filling each section from the evidence package. "
            "If no custom table of contents is supplied, use "
            "`02_document_planning/default_model_methodology_outline.md`.\n"
        ),
        f"{DOCUMENT_TEMPLATE_DIR}/example_table_of_contents.md": (
            "# Example Model Technical Document Table Of Contents\n\n"
            "1. Executive Summary\n"
            "2. Model Purpose, Scope, And Intended Use\n"
            "3. Data, Population, And Target Definition\n"
            "4. Methodology And Model Design\n"
            "5. Variable Selection, Transformations, And Feature Governance\n"
            "6. Model Performance And Validation Evidence\n"
            "7. Sensitivity, Stability, And Limitations\n"
            "8. Implementation, Controls, And Ongoing Use Considerations\n"
            "9. Approvals, Open Items, And Appendices\n"
        ),
    }
    for arcname, text in template_files.items():
        _write_text(
            archive,
            added_arcnames,
            arcname,
            text,
            included,
            evidence_area="document_template",
            source="generated",
            llm_use="Template placeholder for institution-specific document structure.",
        )


def _build_context_markdown(payload: Mapping[str, Any]) -> str:
    run = _mapping(payload.get("run"))
    decision = _mapping(payload.get("decision_summary"))
    lines = [
        "# LLM Model Document Context",
        "",
        "This file is a compact, LLM-readable summary of the completed Quant Studio run.",
        "Use source files listed in `source_citation_map.csv` for citations.",
        "",
        "## Run Overview",
        "",
        f"- Run ID: `{run.get('run_id', 'n/a')}`",
        f"- Execution mode: `{run.get('execution_mode', 'n/a')}`",
        f"- Model type: `{run.get('model_type', 'n/a')}`",
        f"- Target mode: `{run.get('target_mode', 'n/a')}`",
        f"- Target column: `{run.get('target_column', 'n/a')}`",
        f"- Labels available: `{run.get('labels_available', False)}`",
        "",
        "## Decision Summary",
        "",
        f"- Recommendation: `{decision.get('recommendation', 'n/a')}`",
        f"- Decision level: `{str(decision.get('level', 'n/a')).replace('_', ' ').title()}`",
        "",
        "### Rationale",
        "",
    ]
    lines.extend(f"- {item}" for item in decision.get("rationale", []) or ["Not available."])
    lines.extend(["", "## Metrics", ""])
    lines.extend(_markdown_table(decision.get("metrics", [])))
    lines.extend(["", "## Decision Issues", ""])
    lines.extend(_markdown_table(decision.get("issues", [])))
    lines.extend(["", "## Top Feature Drivers", ""])
    lines.extend(_markdown_table(decision.get("feature_drivers", [])))
    lines.extend(["", "## Diagnostic Table Inventory", ""])
    lines.extend(_markdown_table(payload.get("diagnostic_table_inventory", [])))
    lines.extend(["", "## Warnings", ""])
    warnings = payload.get("warnings", [])
    lines.extend(f"- {warning}" for warning in warnings or ["No run warnings were captured."])
    return "\n".join(lines).strip() + "\n"


def _build_model_facts_digest(payload: Mapping[str, Any]) -> str:
    run = _mapping(payload.get("run"))
    decision = _mapping(payload.get("decision_summary"))
    feature_importance = _mapping(payload.get("feature_importance")).get("records", [])
    table_inventory = payload.get("diagnostic_table_inventory", [])
    warnings = payload.get("warnings", [])
    metrics = decision.get("metrics", [])
    issues = decision.get("issues", [])

    lines = [
        "# Model Facts Digest",
        "",
        (
            "This digest summarizes high-confidence facts from the completed Quant Studio "
            "run. Use it as orientation, but cite the underlying files listed in "
            "`EVIDENCE_INDEX.csv` and `source_citation_map.csv` for final document claims."
        ),
        "",
        "## Run Identity",
        "",
        f"- Run ID: `{run.get('run_id', 'n/a')}`",
        f"- Execution mode: `{run.get('execution_mode', 'n/a')}`",
        f"- Model type: `{run.get('model_type', 'n/a')}`",
        f"- Target mode: `{run.get('target_mode', 'n/a')}`",
        f"- Target column: `{run.get('target_column', 'n/a')}`",
        f"- Labels available: `{run.get('labels_available', False)}`",
        f"- Input shape: `{_json_safe(run.get('input_shape', {}))}`",
        "",
        "## Decision Summary",
        "",
        f"- Recommendation: `{decision.get('recommendation', 'n/a')}`",
        f"- Decision level: `{decision.get('level', 'n/a')}`",
        "",
        "## Rationale",
        "",
    ]
    lines.extend(f"- {item}" for item in decision.get("rationale", []) or ["Not available."])
    lines.extend(["", "## Key Metric Rows", ""])
    lines.extend(_markdown_table(list(metrics)[:20] if isinstance(metrics, list) else []))
    lines.extend(["", "## Decision Issues", ""])
    lines.extend(_markdown_table(list(issues)[:25] if isinstance(issues, list) else []))
    lines.extend(["", "## Top Feature Drivers", ""])
    lines.extend(
        _markdown_table(
            list(feature_importance)[:20] if isinstance(feature_importance, list) else []
        )
    )
    lines.extend(["", "## Diagnostic Table Inventory", ""])
    table_inventory_rows = list(table_inventory)[:80] if isinstance(table_inventory, list) else []
    lines.extend(_markdown_table(table_inventory_rows))
    lines.extend(["", "## Captured Warnings", ""])
    lines.extend(f"- {warning}" for warning in warnings or ["No run warnings were captured."])
    lines.extend(
        [
            "",
            "## Use Rules",
            "",
            "- Do not cite operator prompt files as model evidence.",
            "- Do not treat this digest as approval or validation sign-off.",
            (
                "- If a claim is not supported by a citable evidence file, mark it "
                "as missing evidence."
            ),
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _build_start_here() -> str:
    return f"""# Start Here

This package is organized to support LLM-assisted model methodology drafting
while keeping evidence, prompts, controls, visuals, and generated outputs
separate.

Recommended workflow:

1. Read `README_LLM_PACKAGE.md`.
2. Review `{START_HERE_DIR}/PACKAGE_MAP.md` to understand the folder layout.
3. Drop your institution-specific table of contents into `{DOCUMENT_TEMPLATE_DIR}/`.
4. Ask the LLM to create `{GENERATED_OUTPUTS_DIR}/resolved_toc_binding.json`.
5. Ask the LLM to create `{GENERATED_OUTPUTS_DIR}/methodology_doc_spec.json`.
6. Run `{TOOLS_DIR}/validate_docx_spec.py`.
7. Run `{TOOLS_DIR}/build_docx_from_spec.py`.
8. Run `{TOOLS_DIR}/validate_docx_output.py`.

Do not ask the LLM to place images and tables free-form. The recommended path
is: custom TOC -> resolved TOC binding -> document spec -> deterministic DOCX
builder -> validation.
"""


def _build_package_map() -> str:
    return f"""# Package Map

| Folder | Purpose |
| --- | --- |
| `{START_HERE_DIR}/` | Package map, model facts digest, evidence checklist, and build profile. |
| `{DOCUMENT_TEMPLATE_DIR}/` | Place your custom table of contents here. |
| `{DOCUMENT_PLANNING_DIR}/` | Section maps, template binding, schemas, and section briefs. |
| `{DOCX_WORKFLOW_DIR}/` | DOCX instructions, style guide, specs, tables, and checklist. |
| `{EVIDENCE_DIR}/` | Citable facts, context, claims, tables, and source artifacts. |
| `{VISUAL_ASSETS_DIR}/` | Figure placement manifest, chart assets, and document-ready PNG subset. |
| `{VALIDATION_CONTROLS_DIR}/` | Citation rules, validation rules, rubrics, and review controls. |
| `{OPERATOR_INSTRUCTIONS_DIR}/` | Copy/paste prompts and tone profiles. Do not cite. |
| `{TOOLS_DIR}/` | Helper scripts for validation and deterministic DOCX creation. |
| `{GENERATED_OUTPUTS_DIR}/` | Workspace for TOC bindings, document specs, DOCX, and notes. |

Root-level files:

- `START_HERE.md`: shortest workflow orientation.
- `README_LLM_PACKAGE.md`: detailed package instructions.
- `EVIDENCE_INDEX.csv`: authoritative citation boundary.
- `DO_NOT_CITE.md`: files that must not support factual model claims.
- `source_citation_map.csv`: citable source index.
- `llm_evidence_manifest.json`: complete included/skipped file manifest.
"""


def _build_readme() -> str:
    return """# LLM Documentation Package

This zip contains curated run evidence for drafting a model methodology or model
technical document with an LLM. It is designed to support model-risk documentation
aligned with guidance such as SR 11-7 and OCC model-risk expectations, but it does
not certify compliance.

Recommended use:

1. Review `README_LLM_PACKAGE.md`, `EVIDENCE_INDEX.csv`, `DO_NOT_CITE.md`,
   `MODEL_FACTS_DIGEST.md`, and `source_citation_map.csv`.
2. Review the organized folders in `00_START_HERE/PACKAGE_MAP.md`.
3. Review `03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md`,
   `03_docx_workflow/MODEL_DOCUMENT_STYLE_GUIDE.md`,
   `03_docx_workflow/DOCX_QUALITY_CHECKLIST.md`,
   `05_visual_assets/figure_placement_manifest.csv`,
   `03_docx_workflow/table_placement_manifest.csv`,
   `02_document_planning/document_section_evidence_map.csv`,
   `04_evidence/claims/approved_claims.json`,
   `02_document_planning/target_document_schema.json`,
   `04_evidence/gaps/documentation_gaps.md`,
   `06_validation_controls/evidence_strength_policy.json`,
   `06_validation_controls/document_completion_rules.json`,
   and `06_validation_controls/citation_rules.md`.
4. Review `00_START_HERE/llm_package_build_profile.json` if package preparation
   took longer than expected.
5. Add your institution-specific table of contents to `01_document_template/` if needed.
6. Use the controlled prompt sequence in `07_operator_prompts/` when you want
   a plan, draft, and validation sequence.
7. Require the LLM to cite source files for factual claims.
8. Run or follow `08_tools/validate_llm_draft.py` after drafting.
9. Have model-development, validation, and governance reviewers verify the draft.

Citation boundary:

- Files marked `cite_as_evidence = true` in `EVIDENCE_INDEX.csv` are available
  as factual evidence.
- Files in `07_operator_prompts/`, `01_document_template/`, `08_tools/`, prompt
  files, rubrics, validators, and package-control files guide behavior only.
- Do not cite files listed in `DO_NOT_CITE.md` as model evidence.

Privacy posture:

- Row-level input data is excluded by default.
- Row-level prediction output is excluded by default.
- Serialized model binaries are excluded by default.
- Full code snapshots and monitoring handoff folders are excluded by default.
- Generated chart evidence is capped. HTML files support review, and a small
  document-ready PNG subset is included for DOCX figure insertion. Use the
  separate Step 5 `Download Individual Images` action when every standalone PNG
  chart file is required.
"""


def _build_prompt_template() -> str:
    return """# Prompt: Generate Model Methodology Document

You are drafting a model methodology / technical model document for a regulated
financial-services model-risk review. Use only the evidence in this package.
Do not cite files in `07_operator_prompts/`, `01_document_template/`, `08_tools/`,
or files listed in `DO_NOT_CITE.md` as model evidence.

Instructions:

1. Review `EVIDENCE_INDEX.csv` and use only files marked `cite_as_evidence = true`
   for factual claims.
2. If `01_document_template/DROP_TABLE_OF_CONTENTS_HERE.md` has been replaced with a
   custom table of contents, follow that structure exactly.
3. If no custom table of contents is supplied, use
   `02_document_planning/default_model_methodology_outline.md`.
4. Cite the source package file for every factual claim. Prefer citations from
   `source_citation_map.csv`, `04_evidence/claims/approved_claims.json`,
   `02_document_planning/document_section_evidence_map.csv`,
   `04_evidence/context/model_document_context.json`,
   `reports/model_development_dossier.md`, `config/run_config.json`,
   `metadata/reproducibility_manifest.json`, and governance tables.
5. Do not invent results, thresholds, approvals, owners, or limitations.
6. If evidence is missing, write `Evidence not found in package` and identify the
   missing source.
7. Use `04_evidence/gaps/documentation_gaps.md` to preserve missing evidence and open review items.
8. Treat the output as a draft requiring qualified human review. Do not state that
   the model is approved, valid, or compliant unless explicit approval evidence is
   present in the package.
9. Follow `02_document_planning/target_document_schema.json`,
   `06_validation_controls/evidence_strength_policy.json`,
   `06_validation_controls/document_completion_rules.json`,
   `06_validation_controls/controlled_vocabulary.json`, and
   `06_validation_controls/regulatory_language_guardrails.md`.
10. If your environment supports file generation, follow
    `03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md` and create
    `09_generated_outputs/model_methodology.docx`. If not, create Markdown with
    image placeholders from `05_visual_assets/figure_placement_manifest.csv`.

Deliverable:

- A complete DOCX model methodology document when possible, otherwise a
  Markdown fallback with explicit figure placeholders.
- A short list of unresolved documentation gaps.
- A citation appendix mapping major sections to package source files.
"""


def _build_docx_prompt_template() -> str:
    return """# Prompt: Create DOCX Model Methodology Document

You are creating a model methodology / technical model document for regulated
model-risk review. Use only the files in the current extracted LLM documentation
package directory.

Primary deliverable:

- If your environment supports Word document generation, create
  `model_methodology.docx`.
- If your environment cannot generate `.docx`, create
  `model_methodology.md` and include explicit image placeholders that a reviewer
  can insert into Word.

Required files to read first:

- `README_LLM_PACKAGE.md`
- `00_START_HERE/PACKAGE_MAP.md`
- `03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md`
- `03_docx_workflow/MODEL_DOCUMENT_STYLE_GUIDE.md`
- `03_docx_workflow/DOCX_QUALITY_CHECKLIST.md`
- `05_visual_assets/figure_placement_manifest.csv`
- `03_docx_workflow/table_placement_manifest.csv`
- `EVIDENCE_INDEX.csv`
- `DO_NOT_CITE.md`
- `source_citation_map.csv`
- `02_document_planning/document_section_evidence_map.csv`
- `02_document_planning/target_document_schema.json`
- `04_evidence/claims/approved_claims.json`
- `04_evidence/gaps/documentation_gaps.md`
- `00_START_HERE/MODEL_FACTS_DIGEST.md`

DOCX rules:

1. Use Word heading styles for all sections.
2. Use numbered sections aligned to the supplied table of contents or default
   outline.
3. Convert important CSV/JSON summaries into clean Word tables, not pasted raw
   text dumps.
4. Insert only figures marked `insert_or_appendix = body` in
   `05_visual_assets/figure_placement_manifest.csv` into the main body.
5. Put secondary figures marked `appendix` in an appendix or omit them if they do
   not support a specific claim.
6. Every figure must have a title, caption, figure number, and package citation.
7. Every factual claim must cite a file marked `cite_as_evidence=true` in
   `EVIDENCE_INDEX.csv`.
8. Do not cite `07_operator_prompts/`, `01_document_template/`, `08_tools/`,
   prompts, rubrics, validators, style guides, or files listed in
   `DO_NOT_CITE.md` as model evidence.
9. Preserve warnings, failed checks, limitations, and documentation gaps.
10. Do not state that the model is approved, validated, compliant, or
    production-ready unless explicit package evidence supports that statement.

Output requirements:

- Create `model_methodology.docx` when possible.
- Include a short `09_generated_outputs/DOCX_BUILD_NOTES.md` explaining which
  images and evidence sources were used.
- If `.docx` cannot be created, provide Markdown with image placeholders using
  the exact package paths from `05_visual_assets/figure_placement_manifest.csv`.
"""


def _build_docx_build_instructions() -> str:
    return """# DOCX Build Instructions

Use this file when asking an LLM or code-enabled assistant to create a Word
document from this package.

## Required Output

- Preferred file: `09_generated_outputs/model_methodology.docx`
- Fallback file: `09_generated_outputs/model_methodology.md` with image placeholders and captions
- Supporting note: `09_generated_outputs/DOCX_BUILD_NOTES.md`

## Document Structure

1. Use the institution-specific table of contents in `01_document_template/` if one
   was supplied.
2. If no custom table of contents exists, use
   `02_document_planning/default_model_methodology_outline.md`.
3. Use the section-by-section evidence folders in
   `02_document_planning/document_sections/` to keep each section tied to
   approved package evidence.
4. Use `02_document_planning/target_document_schema.json` for required sections,
   minimum evidence, and missing-evidence behavior.

## Figure Use

- Use `05_visual_assets/figure_placement_manifest.csv` as the authority for
  which visuals belong in the main document versus the appendix.
- Prefer PNG files from `05_visual_assets/figures/png/` for DOCX insertion.
- HTML chart files are review aids and should not be embedded directly in Word.
- Do not insert every chart. Insert only visuals that support a specific
  methodology, performance, calibration, stability, explainability, or
  limitation claim.
- Every inserted figure must include:
  - figure number
  - descriptive title
  - one- to three-sentence caption
  - package citation
  - short explanation of why the visual matters

## Tables

- Convert high-value CSV/JSON evidence into readable Word tables.
- Keep long evidence inventories in appendices.
- Do not paste raw JSON blocks into the main body unless the document section is
  explicitly technical/reproducibility-focused.

## Citation Controls

- Cite every factual claim using package paths and fields.
- Use `EVIDENCE_INDEX.csv` and `source_citation_map.csv` to confirm whether a
  source is citable.
- Do not cite prompts, validators, style guides, checklists, or files listed in
  `DO_NOT_CITE.md` as model-run evidence.
- If evidence is absent, write `Evidence not found in package`.

## Fallback When DOCX Is Not Supported

If the LLM environment cannot create `.docx`, produce Markdown with this image
placeholder pattern:

```markdown
![Figure 1: Calibration Curve](05_visual_assets/figures/png/calibration_curve.png)

Figure 1. Calibration Curve. Caption text. [source: package/path > field_or_section]
```
"""


def _build_model_document_style_guide() -> str:
    return """# Model Document Style Guide

## Tone

- Use formal model-risk documentation language.
- Prefer precise, evidence-backed statements over promotional language.
- Use `supports`, `indicates`, or `is consistent with` rather than `proves`.
- Do not imply approval, validation sign-off, compliance, or production readiness
  unless explicit package evidence supports it.

## Layout

- Use numbered headings.
- Keep paragraphs short: usually three to six sentences.
- Use tables for metrics, split summaries, top feature drivers, documentation
  gaps, and validation findings.
- Use figures sparingly in the main body and move supporting visuals to the
  appendix.
- Avoid long raw artifact dumps in the body.

## Figure Formatting

- Figure title: bold, descriptive, and tied to the model evidence.
- Caption: explain what the chart shows, why it matters, and the source.
- Use one large figure per page when possible.
- Standard body figure width: 6.5 inches.
- Appendix figure width: 7.0 inches.

## Table Formatting

- Use short table titles.
- Include only fields needed for the document claim.
- Round numeric values consistently where the source does not require exact
  precision.
- Keep full inventories in appendices.

## Required Evidence Language

- Carry forward warnings, failed checks, and limitations.
- Use `Evidence not found in package` for missing facts.
- Label LLM-created language as draft material requiring human review.
"""


def _build_docx_quality_checklist() -> str:
    return """# DOCX Quality Checklist

Before treating an LLM-generated Word document as ready for human review, verify:

- Every factual claim has a package citation.
- Every citation maps to a citable source in `EVIDENCE_INDEX.csv` or
  `source_citation_map.csv`.
- No prompt, style, validator, checklist, or `DO_NOT_CITE.md` file is cited as
  model evidence.
- Every inserted figure appears in `05_visual_assets/figure_placement_manifest.csv`.
- Every inserted figure has a title, figure number, caption, and citation.
- Every table has a title and clear source reference.
- Warnings, failed checks, limitations, and documentation gaps are preserved.
- No raw row-level input data or row-level predictions are included.
- The document does not claim approval, validation sign-off, compliance, or
  production readiness unless explicit evidence supports that statement.
- The document includes a final unresolved-items section.
- A qualified model-development or validation reviewer has reviewed the draft.
"""


def _build_resolved_toc_binding_schema() -> dict[str, Any]:
    return {
        "$schema": "https://json-schema.org/draft/2020-12/schema",
        "title": "Resolved TOC Binding",
        "type": "object",
        "required": ["toc_source", "sections"],
        "properties": {
            "toc_source": {"type": "string"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": [
                        "toc_heading",
                        "canonical_section",
                        "section_order",
                        "draft_status",
                    ],
                    "properties": {
                        "section_order": {"type": "integer"},
                        "toc_heading": {"type": "string"},
                        "canonical_section": {"type": "string"},
                        "evidence_files": {"type": "array", "items": {"type": "string"}},
                        "allowed_figure_ids": {"type": "array", "items": {"type": "string"}},
                        "allowed_table_ids": {"type": "array", "items": {"type": "string"}},
                        "approved_claim_ids": {"type": "array", "items": {"type": "string"}},
                        "missing_evidence": {"type": "array", "items": {"type": "string"}},
                        "draft_status": {
                            "type": "string",
                            "enum": ["ready", "partial", "blocked", "not_applicable"],
                        },
                    },
                    "additionalProperties": True,
                },
            },
        },
        "additionalProperties": True,
    }


def _build_methodology_doc_spec_schema() -> dict[str, Any]:
    return {
        "$schema": "https://json-schema.org/draft/2020-12/schema",
        "title": "Methodology Document Spec",
        "type": "object",
        "required": ["document_title", "sections"],
        "properties": {
            "document_title": {"type": "string"},
            "document_subtitle": {"type": "string"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["heading", "order", "blocks"],
                    "properties": {
                        "order": {"type": "integer"},
                        "heading": {"type": "string"},
                        "canonical_section": {"type": "string"},
                        "blocks": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "required": ["type"],
                                "properties": {
                                    "type": {
                                        "type": "string",
                                        "enum": ["paragraph", "bullet_list", "table", "figure"],
                                    },
                                    "text": {"type": "string"},
                                    "items": {"type": "array", "items": {"type": "string"}},
                                    "table_id": {"type": "string"},
                                    "figure_id": {"type": "string"},
                                    "caption": {"type": "string"},
                                    "citations": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                    },
                                },
                                "additionalProperties": True,
                            },
                        },
                    },
                    "additionalProperties": True,
                },
            },
            "appendices": {"type": "array"},
        },
        "additionalProperties": True,
    }


def _build_methodology_doc_spec_example() -> dict[str, Any]:
    return {
        "document_title": "Model Methodology Document",
        "document_subtitle": "Generated from Quant Studio LLM package evidence",
        "sections": [
            {
                "order": 1,
                "heading": "Executive Summary",
                "canonical_section": "Executive Summary",
                "blocks": [
                    {
                        "type": "paragraph",
                        "text": (
                            "Draft paragraph with a citation. Replace this with evidence-backed "
                            "content from approved_claims and source files."
                        ),
                        "citations": ["04_evidence/claims/approved_claims.json"],
                    }
                ],
            },
            {
                "order": 2,
                "heading": "Model Performance",
                "canonical_section": (
                    "Performance, Calibration, Stability, And Statistical Testing"
                ),
                "blocks": [
                    {
                        "type": "table",
                        "table_id": "TABLE_001_EXAMPLE",
                        "caption": "Example table placeholder.",
                        "citations": ["03_docx_workflow/table_placement_manifest.csv"],
                    },
                    {
                        "type": "figure",
                        "figure_id": "FIG_001_EXAMPLE",
                        "caption": "Example figure placeholder.",
                        "citations": ["05_visual_assets/figure_placement_manifest.csv"],
                    },
                ],
            },
        ],
        "appendices": [],
    }


def _build_generated_outputs_readme() -> str:
    return """# Generated Outputs Workspace

Use this folder for files created by the LLM and helper scripts.

Expected files:

- `resolved_toc_binding.json`: LLM-created mapping from your custom table of
  contents to package evidence sections.
- `methodology_doc_spec.json`: LLM-created structured document spec.
- `model_methodology.docx`: DOCX created by `08_tools/build_docx_from_spec.py`.
- `model_methodology.md`: fallback when DOCX generation is not available.
- `DOCX_BUILD_NOTES.md`: notes about inserted/omitted figures, tables, and
  unresolved formatting constraints.

Recommended order:

1. Ask the LLM to create `resolved_toc_binding.json`.
2. Ask the LLM to create `methodology_doc_spec.json`.
3. Run `python ../08_tools/validate_docx_spec.py`.
4. Run `python ../08_tools/build_docx_from_spec.py`.
5. Run `python ../08_tools/validate_docx_output.py`.
"""


def _build_resolved_toc_binding_instructions() -> str:
    return f"""# Resolved TOC Binding Instructions

Use this file when an institution-specific table of contents is dropped into
`{DOCUMENT_TEMPLATE_DIR}/`.

The LLM should not immediately draft the methodology document. It should first
create `{GENERATED_OUTPUTS_DIR}/resolved_toc_binding.json` that maps the custom
TOC to package evidence.

Required inputs:

- `{DOCUMENT_TEMPLATE_DIR}/DROP_TABLE_OF_CONTENTS_HERE.md` or another custom TOC
  file in the same folder.
- `{DOCUMENT_PLANNING_DIR}/document_section_evidence_map.csv`
- `{DOCUMENT_PLANNING_DIR}/template_binding.json`
- `{DOCUMENT_PLANNING_DIR}/target_document_schema.json`
- `{VISUAL_ASSETS_DIR}/figure_placement_manifest.csv`
- `{DOCX_WORKFLOW_DIR}/table_placement_manifest.csv`
- `{EVIDENCE_DIR}/claims/approved_claims.json`
- `{EVIDENCE_DIR}/gaps/documentation_gaps.md`
- `EVIDENCE_INDEX.csv`

Required output:

- `{GENERATED_OUTPUTS_DIR}/resolved_toc_binding.json`

Binding rules:

- Preserve the custom TOC heading text and order.
- Map every custom heading to the closest canonical Quant Studio section.
- List the exact evidence files, approved claim IDs, figure IDs, and table IDs
  allowed for that section.
- Mark sections as `blocked` when evidence is missing instead of inventing
  content.
- Do not cite prompt, template, validation, or tooling files as model evidence.
"""


def _build_three_prompt_text() -> str:
    return """Quant Studio LLM Documentation Package Prompt Sequence
======================================================

Use these prompts in order after downloading and extracting the Step 5
Download LLM Package zip.

Prompt 1: Create The Documentation Plan
---------------------------------------

You are assisting with a regulated model-risk technical documentation review.
Use only the files in the current extracted LLM documentation package directory.
Do not cite or rely on files in 07_operator_prompts/ as model evidence. Those
files are workflow instructions only.

First, do not draft the model methodology document. Review the package and
produce a controlled documentation plan.

Read these files first:
- README_LLM_PACKAGE.md
- 00_START_HERE/PACKAGE_MAP.md
- 03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md
- 03_docx_workflow/MODEL_DOCUMENT_STYLE_GUIDE.md
- EVIDENCE_INDEX.csv
- DO_NOT_CITE.md
- 00_START_HERE/MODEL_FACTS_DIGEST.md
- llm_evidence_manifest.json
- source_citation_map.csv
- 05_visual_assets/figure_placement_manifest.csv
- 03_docx_workflow/table_placement_manifest.csv
- 02_document_planning/target_document_schema.json
- 02_document_planning/template_binding.json
- 02_document_planning/document_section_evidence_map.csv
- 04_evidence/claims/approved_claims.json
- 04_evidence/gaps/documentation_gaps.md
- 06_validation_controls/evidence_strength_policy.json
- 06_validation_controls/document_completion_rules.json
- 06_validation_controls/controlled_vocabulary.json
- 06_validation_controls/regulatory_language_guardrails.md

For each required document section, return:
- the section heading to use
- the writing objective
- the citable evidence files and fields that support the section
- approved claims that can be used
- missing evidence or documentation gaps
- high-risk claims that require human review
- whether the section is ready to draft, partially ready, or blocked

Rules:
- Use only files marked cite_as_evidence=true in EVIDENCE_INDEX.csv for factual
  claims.
- Do not cite files listed in DO_NOT_CITE.md.
- Cite package evidence for every factual statement using this style:
  [source: package/path > field_or_section]
- Do not invent results, thresholds, owners, approvals, limitations, or
  regulatory conclusions.
- Do not state that the model is approved, validated, compliant, or
  production-ready unless explicit approval evidence exists in the package.
- If evidence is missing, write `Evidence not found in package`.
- Preserve all warnings, failed checks, documentation gaps, and limitations.
- Treat your output as a draft planning aid that requires qualified human review.

Deliver only:
1. A table with these columns: section, objective, citable_evidence,
   approved_claims, missing_evidence, human_review_needed, draft_status.
2. A list of missing or weak evidence.
3. A list of files that must not be cited.
4. A list of questions for the model owner or validation reviewer.
5. A recommendation on whether to proceed to drafting the full methodology
   document.

Prompt 2: Draft From The Approved Plan
--------------------------------------

Use this only after a qualified reviewer has reviewed the plan from Prompt 1.

You are assisting with a regulated model-risk technical documentation draft.
Use only the files in the current extracted LLM documentation package directory
and the reviewed documentation plan from the prior step.
Use only files marked cite_as_evidence=true in EVIDENCE_INDEX.csv for factual
claims. Do not cite 07_operator_prompts/, 01_document_template/, 08_tools/, prompt
files, rubrics, validators, or files listed in DO_NOT_CITE.md as model evidence.

Draft the full model methodology / technical model document. If your
environment supports Word file generation, create `model_methodology.docx`.
If your environment cannot create `.docx`, create Markdown with explicit image
placeholders using paths from `05_visual_assets/figure_placement_manifest.csv`.

Required sources:
- the reviewed documentation plan
- 03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md
- 03_docx_workflow/MODEL_DOCUMENT_STYLE_GUIDE.md
- 03_docx_workflow/DOCX_QUALITY_CHECKLIST.md
- 05_visual_assets/figure_placement_manifest.csv
- 03_docx_workflow/table_placement_manifest.csv
- EVIDENCE_INDEX.csv
- DO_NOT_CITE.md
- 00_START_HERE/MODEL_FACTS_DIGEST.md
- 02_document_planning/target_document_schema.json
- 02_document_planning/template_binding.json
- 02_document_planning/default_model_methodology_outline.md, unless a custom
  table of contents was provided in 01_document_template/
- 04_evidence/claims/approved_claims.json
- 02_document_planning/document_section_evidence_map.csv
- 04_evidence/context/model_document_context.json
- 04_evidence/interpretation_briefs/metrics_interpretation_brief.md
- 04_evidence/interpretation_briefs/feature_dictionary_narrative.md
- 04_evidence/interpretation_briefs/chart_interpretation_brief.md
- 04_evidence/gaps/documentation_gaps.md
- 06_validation_controls/evidence_strength_policy.json
- 06_validation_controls/controlled_vocabulary.json
- 06_validation_controls/regulatory_language_guardrails.md
- source_citation_map.csv

Rules:
- Draft only from the reviewed documentation plan and citable evidence files.
- Use 04_evidence/claims/approved_claims.json as the primary claim library.
- Follow 02_document_planning/target_document_schema.json for required sections.
- Use the custom table of contents if one was supplied; otherwise use
  02_document_planning/default_model_methodology_outline.md.
- Cite every factual claim using this style:
  [source: package/path > field_or_section]
- When chart image files are present in 05_visual_assets/figures/png/, insert
  only the body-priority visuals identified in
  05_visual_assets/figure_placement_manifest.csv.
  If creating Markdown fallback output, use this form:
  ![Descriptive chart title](05_visual_assets/figures/png/chart.png)
- For every inserted chart, include a short caption explaining what the chart
  shows, why it matters, and the package source citation.
- Do not embed every available chart. Prefer charts that support key
  model-development, validation, calibration, stability, explainability, or
  limitation conclusions.
- Preserve unresolved documentation gaps, warnings, failed checks, and
  limitations.
- Do not invent missing evidence.
- Do not claim approval, validation sign-off, compliance, or production
  readiness unless explicit package evidence supports that statement.
- Use conservative model-risk language from
  06_validation_controls/controlled_vocabulary.json and
  06_validation_controls/regulatory_language_guardrails.md.

Deliver:
1. A complete `model_methodology.docx` file when possible, otherwise a Markdown
   fallback with figure placeholders.
2. A documentation gaps section.
3. A limitations and assumptions section.
4. A citation appendix mapping major sections to package evidence.
5. A short list of items requiring human review before the document can be used.
6. `09_generated_outputs/DOCX_BUILD_NOTES.md` summarizing inserted figures,
   omitted figures, and any formatting limitations.

Prompt 3: Validate The Draft Against Evidence
---------------------------------------------

Use this after the LLM has produced the draft.

You are acting as a model validation and documentation quality reviewer.
Use only the current extracted LLM documentation package directory and the draft
model methodology document.
Do not treat 07_operator_prompts/, 01_document_template/, 08_tools/, prompt files,
rubrics, validators, or files listed in DO_NOT_CITE.md as factual model
evidence.

Review the draft against:
- EVIDENCE_INDEX.csv
- DO_NOT_CITE.md
- 06_validation_controls/document_completion_rules.json
- 06_validation_controls/draft_validation_rules.json
- 06_validation_controls/document_quality_rubric.md
- 03_docx_workflow/DOCX_QUALITY_CHECKLIST.md
- 06_validation_controls/citation_coverage_validator.md
- 06_validation_controls/unsupported_claim_detector.md
- 06_validation_controls/evidence_strength_policy.json
- 06_validation_controls/controlled_vocabulary.json
- 06_validation_controls/regulatory_language_guardrails.md
- source_citation_map.csv
- 04_evidence/gaps/documentation_gaps.md
- 04_evidence/claims/approved_claims.json

Perform these checks:
- Identify factual claims without citations.
- Identify citations that do not map to package evidence.
- Identify unsupported claims, especially approval, validation, compliance, or
  production-readiness claims.
- Identify quantitative values that should be verified against source evidence.
- Confirm that warnings, failed checks, limitations, and documentation gaps were
  preserved.
- Confirm every inserted figure has a title, caption, citation, and appears in
  05_visual_assets/figure_placement_manifest.csv.
- Confirm the DOCX or Markdown fallback follows
  03_docx_workflow/MODEL_DOCUMENT_STYLE_GUIDE.md.
- Score the draft using 06_validation_controls/document_quality_rubric.md.
- Identify sections that are complete, partially supported, missing evidence, or
  blocked.

Rules:
- Do not rewrite the full document.
- Do not approve the model or the document.
- If evidence is missing, write `Evidence not found in package`.
- Treat this as a human-review aid, not a final validation sign-off.

Deliver:
1. Citation coverage findings.
2. Unsupported or high-risk claim findings.
3. Missing evidence and documentation-gap findings.
4. Rubric score by category.
5. Required revisions before human review.
6. A final status of `ready_for_human_review`, `needs_revision`, or
   `blocked_by_missing_evidence`.
"""


def _build_default_outline() -> str:
    return """# Default Model Methodology Outline

## 1. Executive Summary

## 2. Model Purpose, Scope, And Intended Use

## 3. Portfolio, Population, Target, And Performance Horizon

## 4. Data Sources, Preparation, And Split Design

## 5. Model Methodology And Candidate Model Rationale

## 6. Feature Selection, Transformations, Imputation, And Lineage

## 7. Model Performance, Calibration, Stability, And Statistical Testing

## 8. Sensitivity, Scenario, Explainability, And Limitations

## 9. Implementation, Reproducibility, And Controls

## 10. Validation Evidence, Open Items, And Approval Considerations

## 11. Appendices And Source Evidence Index
"""


def _build_section_evidence_map(payload: Mapping[str, Any]) -> list[dict[str, Any]]:
    return [
        _section_row(
            payload,
            section="Executive Summary",
            objective=(
                "Summarize the model purpose, run outcome, primary metrics, key risks, "
                "and remaining review items."
            ),
            sources=[
                "model_document_context.md",
                "reports/decision_summary.md",
                "approved_claims.json",
                "documentation_gaps.md",
            ],
            fields=[
                "run.model_type",
                "run.target_mode",
                "decision_summary.recommendation",
                "decision_summary.metrics",
                "decision_summary.issues",
            ],
            instruction=(
                "Use concise factual language. Do not state that the model is approved "
                "unless explicit approval evidence is present."
            ),
        ),
        _section_row(
            payload,
            section="Purpose, Scope, And Intended Use",
            objective="Describe what the model is intended to estimate and how it may be used.",
            sources=[
                "reports/model_development_dossier.md",
                "reports/model_documentation_pack.md",
                "config/run_config.json",
            ],
            fields=["run.execution_mode", "run.model_type", "run.target_column"],
            instruction=(
                "If business purpose, owner, portfolio, or use-case fields are missing, "
                "flag them as documentation gaps."
            ),
        ),
        _section_row(
            payload,
            section="Data, Population, Target, And Split Design",
            objective=(
                "Document input population, target definition, performance horizon, "
                "train/validation/test design, and leakage controls."
            ),
            sources=[
                "model_document_context.json",
                "config/run_config.json",
                "metadata/reproducibility_manifest.json",
                "04_evidence/table_previews/split_summary.csv",
            ],
            fields=[
                "run.input_shape",
                "run.split_summary",
                "run.target_column",
                "run.labels_available",
            ],
            instruction=(
                "State the split strategy and cite the config. If dates or custom split "
                "labels are used, explain why that design reduces leakage risk."
            ),
        ),
        _section_row(
            payload,
            section="Methodology And Model Design",
            objective="Explain selected model family, estimation approach, and major settings.",
            sources=[
                "model_type_writing_guide.md",
                "config/run_config.json",
                "code/generated_run.py",
                "reports/model_development_dossier.md",
            ],
            fields=["run.model_type", "run.target_mode", "feature_columns"],
            instruction=(
                "Use the model-specific guide. Explain why the model type is appropriate "
                "for the target and where the evidence is only contextual."
            ),
        ),
        _section_row(
            payload,
            section="Features, Transformations, Imputation, And Lineage",
            objective=(
                "Document selected features, source-to-model lineage, transformations, "
                "imputation, binning, and feature review decisions."
            ),
            sources=[
                "feature_dictionary_narrative.md",
                "model/feature_lineage_map.csv",
                "tables/governance/feature_lineage_map.*",
                "04_evidence/table_previews/",
            ],
            fields=["feature_columns", "feature_importance", "diagnostic_table_previews"],
            instruction=(
                "Use feature lineage for factual statements. Do not infer business meaning "
                "for a feature when the dictionary definition is missing."
            ),
        ),
        _section_row(
            payload,
            section="Performance, Calibration, Stability, And Statistical Testing",
            objective=(
                "Explain performance metrics, calibration, stability, statistical tests, "
                "and model-suitability issues."
            ),
            sources=[
                "metrics_interpretation_brief.md",
                "metadata/metrics.json",
                "metadata/statistical_tests.json",
                "reports/validation_pack.md",
                "chart_interpretation_brief.md",
            ],
            fields=[
                "decision_summary.metrics",
                "decision_summary.issues",
                "statistical_tests",
            ],
            instruction=(
                "Use great/good/watch/bad interpretations where available. If a test "
                "failed or was unavailable, describe the limitation plainly."
            ),
        ),
        _section_row(
            payload,
            section="Assumptions, Limitations, And Open Review Items",
            objective="List assumptions, limitations, failed checks, warnings, and gaps.",
            sources=[
                "documentation_gaps.md",
                "reports/validation_pack.md",
                "reports/model_development_dossier.md",
                "llm_evidence_checklist.md",
            ],
            fields=["warnings", "decision_summary.issues"],
            instruction=(
                "Preserve warnings and failed checks. Do not resolve or dismiss open items "
                "unless the evidence package explicitly documents the disposition."
            ),
        ),
        _section_row(
            payload,
            section="Implementation, Reproducibility, And Controls",
            objective=(
                "Document reproducibility assets, generated code, package versions, "
                "artifact manifest, and implementation controls."
            ),
            sources=[
                "code/generated_run.py",
                "code/HOW_TO_RERUN.md",
                "metadata/reproducibility_manifest.json",
                "artifact_manifest.json",
            ],
            fields=["run.run_id", "artifacts"],
            instruction=(
                "Explain that the GUI is not required to rerun the exported workflow. "
                "Cite the config and generated-run files."
            ),
        ),
        _section_row(
            payload,
            section="Appendices And Source Evidence Index",
            objective="Provide source evidence references and unresolved evidence gaps.",
            sources=[
                "source_citation_map.csv",
                "llm_evidence_manifest.json",
                "documentation_gaps.csv",
                "llm_evidence_checklist.csv",
            ],
            fields=["artifacts", "diagnostic_table_inventory"],
            instruction=(
                "Include the citation appendix and missing evidence register rather than "
                "embedding full raw data or row-level prediction records."
            ),
        ),
    ]


def _section_row(
    payload: Mapping[str, Any],
    *,
    section: str,
    objective: str,
    sources: list[str],
    fields: list[str],
    instruction: str,
) -> dict[str, Any]:
    return {
        "document_section": section,
        "writing_objective": objective,
        "primary_sources": sources,
        "source_fields": fields,
        "evidence_status": _source_status(payload, sources),
        "citation_instruction": instruction,
    }


def _write_section_evidence_folders(
    *,
    archive: ZipFile,
    added_arcnames: set[str],
    section_map: list[dict[str, Any]],
    included: list[dict[str, Any]],
) -> None:
    for index, row in enumerate(section_map, start=1):
        section_name = str(row.get("document_section", f"Section {index}"))
        folder = (
            f"{DOCUMENT_PLANNING_DIR}/document_sections/"
            f"{index:02d}_{_safe_name(section_name)}"
        )
        _write_text(
            archive,
            added_arcnames,
            f"{folder}/SECTION_BRIEF.md",
            _build_section_brief_markdown(index=index, row=row),
            included,
            evidence_area="section_evidence_map",
            source="generated",
            llm_use=(
                "Section-specific drafting brief that maps document content to citable "
                "package evidence."
            ),
        )
        _write_csv(
            archive,
            added_arcnames,
            f"{folder}/section_evidence.csv",
            [_section_evidence_csv_row(row)],
            included,
            evidence_area="section_evidence_map",
            source="generated",
            llm_use="Section-specific source map for DOCX methodology drafting.",
        )


def _build_section_brief_markdown(*, index: int, row: Mapping[str, Any]) -> str:
    section_name = str(row.get("document_section", f"Section {index}"))
    sources = row.get("primary_sources", [])
    fields = row.get("source_fields", [])
    lines = [
        f"# {index:02d}. {section_name}",
        "",
        f"Objective: {row.get('writing_objective', '')}",
        "",
        f"Evidence status: `{row.get('evidence_status', '')}`",
        "",
        "## Primary Sources",
        "",
    ]
    lines.extend(f"- `{source}`" for source in sources if source)
    lines.extend(["", "## Source Fields", ""])
    lines.extend(f"- `{field}`" for field in fields if field)
    lines.extend(
        [
            "",
            "## Writing Instruction",
            "",
            str(row.get("citation_instruction", "")),
            "",
            "## DOCX Guidance",
            "",
            "- Use this brief to draft the corresponding Word section.",
            "- Use citable source files for factual claims.",
            "- Move missing or weak evidence into the open-items section.",
            "- Do not cite this brief as the only evidence for a factual claim when an "
            "underlying source file is available.",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _section_evidence_csv_row(row: Mapping[str, Any]) -> dict[str, Any]:
    return {
        "document_section": row.get("document_section", ""),
        "writing_objective": row.get("writing_objective", ""),
        "primary_sources": "; ".join(str(source) for source in row.get("primary_sources", [])),
        "source_fields": "; ".join(str(field) for field in row.get("source_fields", [])),
        "evidence_status": row.get("evidence_status", ""),
        "citation_instruction": row.get("citation_instruction", ""),
    }


def _build_regulatory_crosswalk(payload: Mapping[str, Any]) -> list[dict[str, Any]]:
    rows = [
        {
            "guidance_area": "Model purpose and intended use",
            "expected_documentation": (
                "State business purpose, model use, portfolio, target, horizon, and "
                "material constraints."
            ),
            "package_source": (
                "reports/model_development_dossier.md; config/run_config.json; "
                "model_document_context.md"
            ),
            "evidence_status": _source_status(
                payload,
                [
                    "reports/model_development_dossier.md",
                    "config/run_config.json",
                    "model_document_context.md",
                ],
            ),
            "recommended_document_language": (
                "The model was developed for the documented target and use case. "
                "Any missing owner, portfolio, or horizon fields should be disclosed."
            ),
        },
        {
            "guidance_area": "Conceptual soundness and methodology",
            "expected_documentation": (
                "Explain model family, target compatibility, assumptions, feature design, "
                "and why the chosen approach is appropriate."
            ),
            "package_source": (
                "model_type_writing_guide.md; config/run_config.json; "
                "code/generated_run.py"
            ),
            "evidence_status": _source_status(
                payload,
                ["model_type_writing_guide.md", "config/run_config.json", "code/generated_run.py"],
            ),
            "recommended_document_language": (
                "Use the selected model-family guide and cite the resolved configuration. "
                "Do not claim conceptual soundness beyond the captured evidence."
            ),
        },
        {
            "guidance_area": "Data quality and representativeness",
            "expected_documentation": (
                "Describe source data, input shape, target availability, missingness, "
                "transformations, and sample/split design."
            ),
            "package_source": (
                "model_document_context.json; metadata/reproducibility_manifest.json; "
                "04_evidence/table_previews/"
            ),
            "evidence_status": _source_status(
                payload,
                ["model_document_context.json", "metadata/reproducibility_manifest.json"],
            ),
            "recommended_document_language": (
                "State what the package proves about the input and avoid drawing "
                "representativeness conclusions when source population evidence is absent."
            ),
        },
        {
            "guidance_area": "Outcomes analysis and performance",
            "expected_documentation": (
                "Document split metrics, calibration, backtesting, statistical tests, "
                "suitability checks, and model limitations."
            ),
            "package_source": (
                "metrics_interpretation_brief.md; metadata/metrics.json; "
                "metadata/statistical_tests.json; reports/validation_pack.md"
            ),
            "evidence_status": _source_status(
                payload,
                [
                    "metrics_interpretation_brief.md",
                    "metadata/metrics.json",
                    "metadata/statistical_tests.json",
                    "reports/validation_pack.md",
                ],
            ),
            "recommended_document_language": (
                "Use metric interpretations and failed-check messages. Explain gaps or "
                "warnings rather than converting them into approvals."
            ),
        },
        {
            "guidance_area": "Implementation verification and reproducibility",
            "expected_documentation": (
                "Show how the run can be reproduced, which code/config was used, and "
                "which artifacts support implementation."
            ),
            "package_source": (
                "code/generated_run.py; code/HOW_TO_RERUN.md; "
                "metadata/reproducibility_manifest.json; artifact_manifest.json"
            ),
            "evidence_status": _source_status(
                payload,
                [
                    "code/generated_run.py",
                    "code/HOW_TO_RERUN.md",
                    "metadata/reproducibility_manifest.json",
                    "artifact_manifest.json",
                ],
            ),
            "recommended_document_language": (
                "Describe the exported config and generated runner as the reproducibility "
                "contract. Identify missing code snapshot evidence if not included."
            ),
        },
        {
            "guidance_area": "Governance, limitations, and human review",
            "expected_documentation": (
                "Document unresolved warnings, evidence gaps, reviewer checks, and "
                "limitations before any model-use decision."
            ),
            "package_source": (
                "documentation_gaps.md; human_review_checklist.md; "
                "llm_evidence_checklist.md; source_citation_map.csv"
            ),
            "evidence_status": _source_status(
                payload,
                [
                    "documentation_gaps.md",
                    "human_review_checklist.md",
                    "llm_evidence_checklist.md",
                    "source_citation_map.csv",
                ],
            ),
            "recommended_document_language": (
                "State that the LLM output is a draft and must be verified by qualified "
                "model-development and validation reviewers."
            ),
        },
    ]
    return rows


def _build_approved_claims(payload: Mapping[str, Any]) -> list[dict[str, Any]]:
    run = _mapping(payload.get("run"))
    decision = _mapping(payload.get("decision_summary"))
    metrics = list(decision.get("metrics", []))
    issues = list(decision.get("issues", []))
    claims = [
        _claim(
            "run_id",
            f"The completed Quant Studio run ID is `{run.get('run_id', 'n/a')}`.",
            "model_document_context.json > run.run_id",
            "high",
        ),
        _claim(
            "execution_mode",
            f"The workflow execution mode was `{run.get('execution_mode', 'n/a')}`.",
            "model_document_context.json > run.execution_mode",
            "high",
        ),
        _claim(
            "model_type",
            f"The selected model type was `{run.get('model_type', 'n/a')}`.",
            "model_document_context.json > run.model_type",
            "high",
        ),
        _claim(
            "target_mode",
            (
                f"The model target mode was `{run.get('target_mode', 'n/a')}` using "
                f"target column `{run.get('target_column', 'n/a')}`."
            ),
            "model_document_context.json > run.target_mode, run.target_column",
            "high",
        ),
        _claim(
            "decision_recommendation",
            (
                f"The automated decision summary recommendation was "
                f"`{decision.get('recommendation', 'n/a')}`."
            ),
            "model_document_context.json > decision_summary.recommendation",
            "high",
            claim_type="interpretation_or_recommendation",
            risk_level="high",
            needs_review=True,
        ),
        _claim(
            "feature_count",
            (
                "The exported run context lists "
                f"`{len(payload.get('feature_columns', []))}` selected feature columns."
            ),
            "model_document_context.json > feature_columns",
            "high",
        ),
    ]
    for metric in metrics[:10]:
        if not isinstance(metric, Mapping):
            continue
        claims.append(
            _claim(
                f"metric_{metric.get('split', 'split')}_{metric.get('metric', 'metric')}",
                (
                    f"Metric `{metric.get('metric', 'n/a')}` on split "
                    f"`{metric.get('split', 'n/a')}` was `{metric.get('value', 'n/a')}` "
                    f"with interpretation `{metric.get('interpretation', 'n/a')}`."
                ),
                "model_document_context.json > decision_summary.metrics",
                "high",
                claim_type="quantitative_result",
                risk_level="medium",
            )
        )
    for issue in issues[:10]:
        if not isinstance(issue, Mapping):
            continue
        severity = str(issue.get("severity", "")).lower()
        if severity and severity != "none":
            claims.append(
                _claim(
                    f"issue_{len(claims)}",
                    (
                        f"The decision summary recorded `{issue.get('severity', 'n/a')}` "
                        f"from `{issue.get('source', 'n/a')}`: "
                        f"{issue.get('message', 'n/a')}"
                    ),
                    "model_document_context.json > decision_summary.issues",
                    "high",
                    claim_type="limitation_or_issue",
                    risk_level="high",
                    needs_review=True,
                )
            )
    warnings = list(payload.get("warnings", []))
    if warnings:
        claims.append(
            _claim(
                "warnings_recorded",
                f"The run recorded `{len(warnings)}` warning(s) that require review.",
                "model_document_context.json > warnings",
                "high",
                claim_type="limitation_or_issue",
                risk_level="high",
                needs_review=True,
            )
        )
    else:
        claims.append(
            _claim(
                "no_warnings_recorded",
                "The run context did not capture pipeline warnings.",
                "model_document_context.json > warnings",
                "high",
            )
        )
    return claims


def _claim(
    claim_id: str,
    claim_text: str,
    source_reference: str,
    confidence: str,
    *,
    claim_type: str = "run_fact",
    risk_level: str = "low",
    needs_review: bool = False,
) -> dict[str, Any]:
    return {
        "claim_id": claim_id,
        "claim_type": claim_type,
        "approved_claim": claim_text,
        "source_reference": source_reference,
        "confidence": confidence,
        "risk_level": "high" if needs_review else risk_level,
        "requires_human_review": needs_review,
        "requires_validator_review": needs_review or risk_level.lower() == "high",
        "llm_instruction": (
            "Use this claim only with the citation shown. Do not expand beyond the evidence."
        ),
    }


def _build_target_document_schema(
    payload: Mapping[str, Any],
    section_map: list[dict[str, Any]],
) -> dict[str, Any]:
    run = _mapping(payload.get("run"))
    sections = []
    for index, row in enumerate(section_map, start=1):
        section = str(row.get("document_section", f"Section {index}"))
        evidence_status = str(row.get("evidence_status", "not_found"))
        sections.append(
            {
                "section_id": f"{index:02d}",
                "canonical_heading": section,
                "required": True,
                "writing_objective": row.get("writing_objective", ""),
                "primary_sources": row.get("primary_sources", []),
                "required_source_fields": row.get("source_fields", []),
                "minimum_evidence_status": (
                    "present" if evidence_status == "present" else "partial_with_gap_note"
                ),
                "minimum_citations": 1,
                "must_include": [
                    "facts explicitly supported by package evidence",
                    "material limitations or missing evidence relevant to this section",
                ],
                "must_not_include": [
                    "approval conclusions unless approval evidence is present",
                    "owner, policy, threshold, or use-case details not found in evidence",
                    "row-level data excerpts or personal data",
                ],
                "when_evidence_missing": (
                    "Write `Evidence not found in package` and add the item to the "
                    "documentation gaps section."
                ),
            }
        )
    return {
        "schema_version": "1.0",
        "package_version": PACKAGE_VERSION,
        "intended_output": (
            "Model methodology / technical model document in DOCX when the LLM "
            "environment supports file generation; Markdown with figure placeholders "
            "is the fallback."
        ),
        "run_context": {
            "run_id": run.get("run_id", ""),
            "model_type": run.get("model_type", ""),
            "target_mode": run.get("target_mode", ""),
            "target_column": run.get("target_column", ""),
        },
        "global_requirements": [
            "Use the supplied table of contents when present; otherwise use the default outline.",
            "Every factual claim must cite package evidence using the required citation style.",
            "Use 04_evidence/claims/approved_claims.json as the primary claim library.",
            (
                "Use 05_visual_assets/figure_placement_manifest.csv for DOCX "
                "or Markdown figure insertion."
            ),
            (
                "Disclose 04_evidence/gaps/documentation_gaps.md items instead "
                "of filling gaps with assumptions."
            ),
            "Do not state the model is approved, validated, compliant, or production-ready "
            "unless explicit package evidence supports that status.",
        ],
        "required_sections": sections,
        "required_appendices": [
            "source evidence index",
            "documentation gaps and unresolved items",
            "human review checklist completion",
        ],
    }


def _build_evidence_strength_policy() -> dict[str, Any]:
    return {
        "policy_version": "1.0",
        "purpose": "Rank evidence sources so LLM-generated claims stay tied to the run record.",
        "evidence_levels": [
            {
                "level": "authoritative",
                "sources": [
                    "04_evidence/source_artifacts/config/run_config.json",
                    "04_evidence/context/model_document_context.json",
                    "04_evidence/source_artifacts/metadata/metrics.json",
                    "04_evidence/source_artifacts/metadata/statistical_tests.json",
                    "04_evidence/source_artifacts/metadata/reproducibility_manifest.json",
                    "04_evidence/source_artifacts/tables/governance/*.csv",
                    "04_evidence/converted_tables/tables/governance/*.csv",
                ],
                "allowed_claims": [
                    "run configuration",
                    "target definition",
                    "metric values",
                    "statistical-test results",
                    "feature lineage",
                    "reproducibility evidence",
                ],
            },
            {
                "level": "derived_summary",
                "sources": [
                    "04_evidence/claims/approved_claims.json",
                    "02_document_planning/document_section_evidence_map.csv",
                    "04_evidence/interpretation_briefs/metrics_interpretation_brief.md",
                    "04_evidence/interpretation_briefs/feature_dictionary_narrative.md",
                    "04_evidence/interpretation_briefs/chart_interpretation_brief.md",
                    "05_visual_assets/figure_placement_manifest.csv",
                ],
                "allowed_claims": [
                    "plain-English interpretations of structured run facts",
                    "section-to-evidence mapping",
                    "drafting guidance",
                    "DOCX figure placement guidance",
                ],
            },
            {
                "level": "narrative_support",
                "sources": [
                    "reports/model_development_dossier.md",
                    "reports/validation_pack.md",
                    "reports/decision_summary.md",
                ],
                "allowed_claims": [
                    "narrative summaries when consistent with structured evidence",
                    "review context",
                ],
            },
            {
                "level": "prompt_or_template",
                "sources": [
                    "07_operator_prompts/*.txt",
                    "07_operator_prompts/prompts/*.md",
                    "07_operator_prompts/tone_profiles/*.md",
                    "01_document_template/*.md",
                ],
                "allowed_claims": ["formatting and writing instructions only"],
            },
        ],
        "claim_rules": [
            {
                "claim_type": "quantitative_result",
                "minimum_level": "authoritative",
                "requires_exact_value": True,
                "requires_citation": True,
                "requires_human_review": False,
            },
            {
                "claim_type": "methodology_description",
                "minimum_level": "authoritative",
                "requires_citation": True,
                "requires_human_review": False,
            },
            {
                "claim_type": "interpretation_or_recommendation",
                "minimum_level": "derived_summary",
                "requires_citation": True,
                "requires_human_review": True,
            },
            {
                "claim_type": "approval_or_compliance",
                "minimum_level": "authoritative",
                "requires_explicit_approval_source": True,
                "requires_citation": True,
                "requires_human_review": True,
            },
        ],
    }


def _build_document_completion_rules() -> dict[str, Any]:
    return {
        "rule_version": "1.0",
        "blocking_failures": [
            "A factual paragraph has no package citation.",
            "A metric, threshold, coefficient, p-value, or count differs from package evidence.",
            "The draft claims approval, validation sign-off, or regulatory compliance without "
            "explicit evidence.",
            "The draft ignores a high-severity decision issue, warning, or documentation gap.",
            "The draft includes raw row-level input data or row-level predictions.",
            "The DOCX includes a figure without a title, caption, citation, or "
            "figure_placement_manifest.csv entry.",
        ],
        "required_artifact_checks": [
            {
                "file": "approved_claims.json",
                "rule": "All material run facts should be traceable to an approved claim or "
                "an authoritative source file.",
            },
            {
                "file": "documentation_gaps.md",
                "rule": "All missing evidence and unresolved issues must be carried into "
                "the draft.",
            },
            {
                "file": "source_citation_map.csv",
                "rule": "All cited package paths must exist in the source citation map or "
                "be generated control files.",
            },
            {
                "file": "human_review_checklist.md",
                "rule": "Human review must be completed before relying on the document.",
            },
            {
                "file": "figure_placement_manifest.csv",
                "rule": "Inserted DOCX figures should be selected from the placement manifest.",
            },
        ],
        "completion_status_values": [
            {
                "status": "complete_for_review",
                "meaning": "The draft is ready for qualified human review, not approval.",
            },
            {
                "status": "incomplete_missing_evidence",
                "meaning": "The draft depends on evidence not found in the package.",
            },
            {
                "status": "failed_control_check",
                "meaning": "The draft violates at least one blocking rule.",
            },
        ],
    }


def _build_controlled_vocabulary() -> dict[str, Any]:
    return {
        "vocabulary_version": "1.0",
        "preferred_terms": {
            "Quant Studio": "Use for the application name.",
            "model methodology document": "Use for the LLM-generated technical document.",
            "decision summary": "Use for the generated run recommendation evidence.",
            "validation evidence": "Use for tests, diagnostics, and reviewer-facing outputs.",
            "documentation gap": "Use for missing or incomplete evidence.",
            "human review required": "Use whenever LLM output affects model-risk conclusions.",
        },
        "restricted_terms": {
            "approved": "Use only with explicit approval evidence.",
            "validated": "Use only with explicit validation sign-off evidence.",
            "compliant": "Use only with explicit compliance determination evidence.",
            "guarantees": "Do not use for model performance or risk outcomes.",
            "proves": "Prefer `supports`, `indicates`, or `is consistent with`.",
            "production-ready": "Use only with implementation approval evidence.",
        },
        "required_replacements": {
            "the model proves": "the evidence indicates",
            "the model is compliant": "the package does not by itself establish compliance",
            "the model is approved": "approval evidence was not found unless cited explicitly",
            "no risk": "no automated issue was identified in the cited evidence",
        },
        "metric_language": {
            "auc": "ranking/discrimination metric; not calibration by itself",
            "ks": "maximum separation between event and non-event cumulative distributions",
            "brier_score": "probability calibration and sharpness metric for binary models",
            "rmse": "average magnitude of continuous prediction error with squared penalty",
            "mae": "average absolute continuous prediction error",
        },
    }


def _build_draft_validation_rules() -> dict[str, Any]:
    return {
        "validator_version": "1.0",
        "citation_pattern": r"\[source: package/path > field_or_section\]",
        "paragraphs_requiring_citation": [
            "paragraphs containing numbers, percentages, coefficients, p-values, or thresholds",
            "paragraphs naming model type, target, split design, features, or data source",
            "paragraphs interpreting performance, calibration, stability, or suitability",
            "paragraphs describing warnings, limitations, approvals, or governance status",
            "figure captions and table captions",
        ],
        "unsupported_claim_indicators": [
            "approved",
            "validated",
            "compliant",
            "production-ready",
            "unbiased",
            "no limitations",
            "guaranteed",
            "best model",
            "materially complete",
        ],
        "required_review_outputs": [
            "citation coverage percentage",
            "unsupported claim list",
            "uncited quantitative fact list",
            "unresolved documentation gaps",
            "human reviewer sign-off checklist",
        ],
        "docx_specific_checks": [
            "Every inserted figure appears in figure_placement_manifest.csv.",
            "Every inserted figure has a title, caption, and package citation.",
            "The main body does not contain appendix-only figures unless a reviewer "
            "requested them.",
            "The document uses section headings from the supplied template or default outline.",
        ],
    }


def _build_template_binding(section_map: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "binding_version": "1.0",
        "purpose": (
            "Map an institution-specific table of contents to Quant Studio evidence. "
            "Fill `custom_heading` when using an external methodology template."
        ),
        "instructions": [
            "Paste or save the external table of contents in 01_document_template/.",
            "Map each custom heading to one canonical section below.",
            "Do not draft sections with `evidence_status` equal to `not_found` without "
            "explicitly disclosing missing evidence.",
        ],
        "section_bindings": [
            {
                "canonical_section": row.get("document_section", ""),
                "custom_heading": "",
                "required": True,
                "evidence_status": row.get("evidence_status", ""),
                "primary_sources": row.get("primary_sources", []),
                "drafting_instruction": row.get("citation_instruction", ""),
            }
            for row in section_map
        ],
    }


def _build_llm_redaction_policy() -> dict[str, Any]:
    return {
        "policy_version": "1.0",
        "default_posture": "minimum necessary evidence",
        "excluded_by_default": [
            "row-level input snapshots",
            "row-level predictions",
            "serialized model binaries",
            "full code snapshots",
            "monitoring handoff bundles",
            "personal identifiers and account-level details",
        ],
        "allowed_by_default": [
            "run configuration",
            "aggregate metrics",
            "statistical-test results",
            "governance tables",
            "feature lineage",
            "generated run script",
            "narrative reports",
            "capped aggregate chart images for DOCX insertion",
        ],
        "conditional_items": [
            {
                "item": "business owner names or approver names",
                "condition": "include only when institution policy permits LLM processing",
            },
            {
                "item": "portfolio names or product names",
                "condition": "redact or generalize if confidential",
            },
            {
                "item": "custom table of contents",
                "condition": "remove confidential policy language if required",
            },
        ],
        "llm_instruction": (
            "Do not ask the user to provide excluded data unless an approved privacy and "
            "model-governance process authorizes that transfer."
        ),
    }


def _build_document_quality_rubric() -> list[dict[str, Any]]:
    return [
        _rubric_row(
            "Evidence grounding",
            25,
            "Every material factual statement cites package evidence.",
            "Uncited quantitative claims or unsupported conclusions.",
        ),
        _rubric_row(
            "Methodology accuracy",
            20,
            "Model type, target, features, split design, and assumptions match the run config.",
            "Confuses target mode, model family, or feature-treatment evidence.",
        ),
        _rubric_row(
            "Performance and limitations",
            20,
            "Metrics, warnings, failed checks, and documentation gaps are complete and balanced.",
            "Turns warnings into approvals or omits material limitations.",
        ),
        _rubric_row(
            "Regulatory documentation coverage",
            15,
            "Covers purpose, conceptual soundness, data, outcomes analysis, controls, and review.",
            "Leaves major SR 11-7/OCC-style documentation areas unaddressed.",
        ),
        _rubric_row(
            "Reproducibility and implementation controls",
            10,
            "Explains run configuration, generated code, artifact manifest, and rerun evidence.",
            "Cannot be traced back to the exported configuration and code.",
        ),
        _rubric_row(
            "Clarity and review readiness",
            10,
            "Uses conservative, precise language and separates evidence from judgment.",
            "Uses promotional language, overstates certainty, or hides unresolved questions.",
        ),
    ]


def _rubric_row(
    category: str,
    points: int,
    full_credit: str,
    failure_signal: str,
) -> dict[str, Any]:
    return {
        "category": category,
        "points": points,
        "full_credit_standard": full_credit,
        "failure_signal": failure_signal,
        "minimum_acceptable_score": int(points * 0.7),
    }


def _build_document_quality_rubric_markdown() -> str:
    rows = _build_document_quality_rubric()
    lines = [
        "# Document Quality Rubric",
        "",
        "Score an LLM-generated methodology draft before relying on it.",
        "",
    ]
    lines.extend(_markdown_table(rows))
    lines.extend(
        [
            "",
            "Passing guidance:",
            "",
            "- A draft below 80 total points should be revised before formal review.",
            "- Any blocking failure in `document_completion_rules.json` overrides the score.",
            "- A qualified reviewer must verify citations and conclusions.",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _build_citation_coverage_validator() -> str:
    return """# Citation Coverage Validator

Use this check after an LLM drafts the methodology document.

Required checks:

- Every paragraph with a model fact, metric, threshold, date, count, coefficient,
  feature, target definition, split design, warning, or recommendation must cite
  a package source.
- Every citation must follow `[source: package/path > field_or_section]`.
- Every cited package path should appear in `source_citation_map.csv` and be
  marked `cite_as_evidence = true` in `EVIDENCE_INDEX.csv`.
- Claims sourced only to prompts, tone profiles, templates, validators, rubrics,
  or files listed in `DO_NOT_CITE.md` are not factual evidence.
- Missing citations must be marked as `Evidence not found in package` before
  human review.

Recommended output:

- citation coverage percentage
- list of uncited paragraphs
- list of citations to missing package paths
- list of citations to instruction-only sources
"""


def _build_unsupported_claim_detector() -> str:
    return """# Unsupported Claim Detector

Screen the LLM-generated draft for claims that often exceed the evidence.

Flag any sentence that says or implies:

- the model is approved, validated, compliant, production-ready, unbiased, or
  limitation-free without explicit source evidence
- the model will perform the same in production as in development
- a metric proves conceptual soundness or regulatory acceptability by itself
- a feature causes an outcome when the package only shows association or model
  importance
- missing documentation is immaterial unless a reviewer explicitly made that
  determination
- all SR 11-7, OCC, CECL, CCAR, or internal policy requirements are satisfied
  unless the package includes explicit evidence

Required handling:

- Replace unsupported claims with conservative, evidence-backed wording.
- Preserve open issues from `documentation_gaps.md`.
- Escalate high-risk claims to human validation review.
"""


def _build_regulatory_language_guardrails() -> str:
    return """# Regulatory Language Guardrails

Use conservative model-risk language.

Preferred:

- `The evidence indicates...`
- `The run output shows...`
- `The package includes evidence for...`
- `The following limitations require review...`
- `Evidence was not found in the package for...`

Avoid unless explicitly supported:

- `approved`
- `validated`
- `compliant`
- `production-ready`
- `no material risk`
- `all requirements are satisfied`
- `the model proves`

Regulatory references such as SR 11-7, OCC guidance, CECL, or CCAR should be
used as documentation context, not as a claim that the model satisfies those
requirements. State that final determinations require institution-specific
review.
"""


def _build_llm_draft_validator_script() -> str:
    return r'''"""Validate an LLM-generated methodology draft against package controls.

Usage:
    python 08_tools/validate_llm_draft.py path/to/draft.md

This lightweight checker is intentionally conservative. It does not prove that
the document is correct; it highlights missing citations and high-risk language
for human review.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

CITATION_RE = re.compile(r"\[source:\s*([^\]>]+)(?:>\s*([^\]]+))?\]", re.IGNORECASE)
QUANT_RE = re.compile(
    r"(\b\d+(?:\.\d+)?%?\b|\bauc\b|\bks\b|\brmse\b|\bmae\b|\bp-value\b|"
    r"\bcoefficient\b|\bthreshold\b|\bfeature\b|\btarget\b|\bsplit\b)",
    re.IGNORECASE,
)
HIGH_RISK_RE = re.compile(
    r"\b(approved|validated|compliant|production-ready|unbiased|guarantee[sd]?|"
    r"no limitations|no material risk|all requirements are satisfied)\b",
    re.IGNORECASE,
)


def _load_manifest(package_root: Path) -> set[str]:
    manifest_path = package_root / "llm_evidence_manifest.json"
    if not manifest_path.exists():
        return set()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    return {
        str(row.get("package_path", "")).lower()
        for row in manifest.get("included_files", [])
        if isinstance(row, dict)
    }


def _paragraphs(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]


def validate(draft_path: Path) -> dict[str, object]:
    text = draft_path.read_text(encoding="utf-8")
    package_root = draft_path.parent
    if package_root.name != "llm_documentation_package":
        candidate = draft_path.parent / "llm_documentation_package"
        if candidate.exists():
            package_root = candidate
    included_paths = _load_manifest(package_root)
    uncited_quantitative_paragraphs = []
    high_risk_claims = []
    missing_citation_paths = []

    for paragraph in _paragraphs(text):
        has_citation = bool(CITATION_RE.search(paragraph))
        if QUANT_RE.search(paragraph) and not has_citation:
            uncited_quantitative_paragraphs.append(paragraph[:300])
        if HIGH_RISK_RE.search(paragraph):
            high_risk_claims.append(paragraph[:300])

    for match in CITATION_RE.finditer(text):
        cited_path = match.group(1).strip().lower()
        citation_found = cited_path in included_paths or any(
            path.endswith("/" + cited_path) or path.endswith(cited_path)
            for path in included_paths
        )
        if included_paths and not citation_found:
            missing_citation_paths.append(cited_path)

    citation_count = len(CITATION_RE.findall(text))
    paragraph_count = max(len(_paragraphs(text)), 1)
    return {
        "draft_path": str(draft_path),
        "paragraph_count": paragraph_count,
        "citation_count": citation_count,
        "citation_density": round(citation_count / paragraph_count, 4),
        "uncited_quantitative_paragraphs": uncited_quantitative_paragraphs,
        "high_risk_claims": high_risk_claims,
        "missing_citation_paths": sorted(set(missing_citation_paths)),
        "status": (
            "needs_review"
            if uncited_quantitative_paragraphs or high_risk_claims or missing_citation_paths
            else "no_automated_issues_detected"
        ),
    }


def main() -> int:
    if len(sys.argv) != 2:
        sys.stdout.write("Usage: python 08_tools/validate_llm_draft.py path/to/draft.md\n")
        return 2
    result = validate(Path(sys.argv[1]))
    sys.stdout.write(json.dumps(result, indent=2) + "\n")
    return 1 if result["status"] == "needs_review" else 0


if __name__ == "__main__":
    raise SystemExit(main())
'''


def _build_docx_spec_validator_script() -> str:
    return r'''"""Validate methodology_doc_spec.json before DOCX creation.

Usage:
    python ../08_tools/validate_docx_spec.py

Run this from 09_generated_outputs/ or from anywhere inside the extracted
llm_documentation_package folder. The checker intentionally uses only the Python
standard library so it can run in constrained review environments.
"""

from __future__ import annotations

import csv
import json
import sys
from pathlib import Path
from typing import Any


def _package_root() -> Path:
    here = Path(__file__).resolve()
    if here.parent.name == "08_tools":
        return here.parents[1]
    for parent in [here.parent, *here.parents]:
        if (parent / "EVIDENCE_INDEX.csv").exists():
            return parent
    return Path.cwd()


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def _boolish(value: str) -> bool:
    return str(value).strip().lower() in {"true", "1", "yes", "y"}


def _path_exists(package_root: Path, package_path: str) -> bool:
    return bool(package_path) and (package_root / package_path).exists()


def validate() -> dict[str, Any]:
    package_root = _package_root()
    generated_root = package_root / "09_generated_outputs"
    spec_path = generated_root / "methodology_doc_spec.json"
    figure_manifest = _read_csv(package_root / "05_visual_assets" / "figure_placement_manifest.csv")
    table_manifest = _read_csv(package_root / "03_docx_workflow" / "table_placement_manifest.csv")
    evidence_index = _read_csv(package_root / "EVIDENCE_INDEX.csv")
    valid_figure_ids = {row.get("figure_id", "") for row in figure_manifest}
    valid_table_ids = {row.get("table_id", "") for row in table_manifest}
    citable_paths = {
        row.get("package_path", "")
        for row in evidence_index
        if _boolish(row.get("cite_as_evidence", ""))
    }
    non_citable_paths = {
        row.get("package_path", "")
        for row in evidence_index
        if not _boolish(row.get("cite_as_evidence", ""))
    }
    errors: list[str] = []
    warnings: list[str] = []

    if not spec_path.exists():
        errors.append(f"Missing {spec_path}")
        return {"status": "failed", "errors": errors, "warnings": warnings}

    spec = _read_json(spec_path)
    if not isinstance(spec, dict):
        errors.append("Spec root must be a JSON object.")
        return {"status": "failed", "errors": errors, "warnings": warnings}
    if not str(spec.get("document_title", "")).strip():
        errors.append("document_title is required.")
    sections = spec.get("sections")
    if not isinstance(sections, list) or not sections:
        errors.append("sections must be a non-empty list.")
        sections = []

    seen_orders: set[int] = set()
    for section_index, section in enumerate(sections, start=1):
        if not isinstance(section, dict):
            errors.append(f"Section {section_index} must be an object.")
            continue
        heading = str(section.get("heading", "")).strip()
        if not heading:
            errors.append(f"Section {section_index} is missing heading.")
        order = section.get("order")
        if not isinstance(order, int):
            errors.append(f"Section {heading or section_index} has non-integer order.")
        elif order in seen_orders:
            errors.append(f"Duplicate section order {order}.")
        else:
            seen_orders.add(order)
        blocks = section.get("blocks")
        if not isinstance(blocks, list) or not blocks:
            warnings.append(f"Section {heading or section_index} has no blocks.")
            continue
        for block_index, block in enumerate(blocks, start=1):
            if not isinstance(block, dict):
                errors.append(f"Section {heading} block {block_index} must be an object.")
                continue
            block_type = block.get("type")
            if block_type not in {"paragraph", "bullet_list", "table", "figure"}:
                errors.append(
                    f"Section {heading} block {block_index} has invalid "
                    f"type {block_type!r}."
                )
            if block_type == "paragraph" and not str(block.get("text", "")).strip():
                warnings.append(f"Section {heading} paragraph block {block_index} is empty.")
            if block_type == "bullet_list" and not isinstance(block.get("items"), list):
                errors.append(
                    f"Section {heading} bullet_list block {block_index} "
                    "needs items list."
                )
            if block_type == "figure":
                figure_id = str(block.get("figure_id", ""))
                if figure_id not in valid_figure_ids:
                    errors.append(f"Section {heading} references unknown figure_id {figure_id!r}.")
            if block_type == "table":
                table_id = str(block.get("table_id", ""))
                if table_id not in valid_table_ids:
                    errors.append(f"Section {heading} references unknown table_id {table_id!r}.")
            for citation in block.get("citations", []) or []:
                citation_path = str(citation).split(">", 1)[0].strip()
                if citation_path in non_citable_paths:
                    errors.append(f"Section {heading} cites non-citable file {citation_path!r}.")
                elif (
                    citable_paths
                    and citation_path not in citable_paths
                    and not _path_exists(package_root, citation_path)
                ):
                    warnings.append(
                        f"Section {heading} cites a path not found in package "
                        f"index: {citation_path!r}."
                    )

    status = "passed" if not errors else "failed"
    result = {
        "status": status,
        "package_root": str(package_root),
        "spec_path": str(spec_path),
        "section_count": len(sections),
        "errors": errors,
        "warnings": warnings,
    }
    (generated_root / "methodology_doc_spec_validation.json").write_text(
        json.dumps(result, indent=2),
        encoding="utf-8",
    )
    return result


def main() -> int:
    result = validate()
    sys.stdout.write(json.dumps(result, indent=2) + "\n")
    return 0 if result["status"] == "passed" else 1


if __name__ == "__main__":
    raise SystemExit(main())
'''


def _build_docx_builder_script() -> str:
    return r'''"""Build model_methodology.docx from methodology_doc_spec.json.

Usage:
    python ../08_tools/build_docx_from_spec.py

The preferred output is 09_generated_outputs/model_methodology.docx. If the
optional python-docx package is unavailable, this script writes a Markdown
fallback to 09_generated_outputs/model_methodology.md and records the limitation
in DOCX_BUILD_NOTES.md.
"""

from __future__ import annotations

import csv
import json
import sys
from pathlib import Path
from typing import Any


def _package_root() -> Path:
    here = Path(__file__).resolve()
    if here.parent.name == "08_tools":
        return here.parents[1]
    for parent in [here.parent, *here.parents]:
        if (parent / "EVIDENCE_INDEX.csv").exists():
            return parent
    return Path.cwd()


PACKAGE_ROOT = _package_root()
GENERATED_ROOT = PACKAGE_ROOT / "09_generated_outputs"
SPEC_PATH = GENERATED_ROOT / "methodology_doc_spec.json"


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def _table_manifest() -> dict[str, dict[str, str]]:
    rows = _read_csv(PACKAGE_ROOT / "03_docx_workflow" / "table_placement_manifest.csv")
    return {row.get("table_id", ""): row for row in rows if row.get("table_id")}


def _figure_manifest() -> dict[str, dict[str, str]]:
    rows = _read_csv(PACKAGE_ROOT / "05_visual_assets" / "figure_placement_manifest.csv")
    return {row.get("figure_id", ""): row for row in rows if row.get("figure_id")}


def _read_table_preview(
    package_path: str,
    *,
    max_rows: int = 12,
) -> tuple[list[str], list[list[str]]]:
    path = PACKAGE_ROOT / package_path
    if not path.exists():
        return ["missing_source"], [[package_path]]
    if path.suffix.lower() == ".csv":
        with path.open("r", encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        if not rows:
            return ["empty"], [[""]]
        return rows[0], rows[1 : max_rows + 1]
    if path.suffix.lower() == ".json":
        payload = _read_json(path)
        if isinstance(payload, list) and payload and isinstance(payload[0], dict):
            columns = sorted({str(key) for row in payload[:max_rows] for key in row})
            records = [
                [str(row.get(column, "")) for column in columns]
                for row in payload[:max_rows]
            ]
            return columns, records
        if isinstance(payload, dict):
            rows = [
                [str(key), json.dumps(value, default=str)[:500]]
                for key, value in list(payload.items())[:max_rows]
            ]
            return ["field", "value"], rows
    return ["unsupported_source"], [[package_path]]


def _citations_text(citations: Any) -> str:
    if not citations:
        return ""
    if not isinstance(citations, list):
        citations = [citations]
    return " ".join(f"[source: {citation}]" for citation in citations)


def _write_markdown(
    spec: dict[str, Any],
    table_rows: dict[str, dict[str, str]],
    figure_rows: dict[str, dict[str, str]],
) -> Path:
    output_path = GENERATED_ROOT / "model_methodology.md"
    lines = [f"# {spec.get('document_title', 'Model Methodology Document')}", ""]
    subtitle = str(spec.get("document_subtitle", "")).strip()
    if subtitle:
        lines.extend([subtitle, ""])
    for section in sorted(spec.get("sections", []), key=lambda item: item.get("order", 9999)):
        lines.extend([f"## {section.get('heading', 'Untitled Section')}", ""])
        for block in section.get("blocks", []):
            block_type = block.get("type")
            if block_type == "paragraph":
                text = str(block.get("text", "")).strip()
                citation_text = _citations_text(block.get("citations"))
                lines.extend([f"{text} {citation_text}".strip(), ""])
            elif block_type == "bullet_list":
                for item in block.get("items", []) or []:
                    lines.append(f"- {item}")
                lines.append("")
            elif block_type == "figure":
                row = figure_rows.get(str(block.get("figure_id", "")), {})
                image_path = row.get("image_path", "")
                title = row.get("figure_name", block.get("figure_id", "Figure"))
                caption = block.get("caption") or row.get("caption", "")
                if image_path:
                    lines.append(f"![{title}]({image_path})")
                figure_citations = _citations_text(
                    block.get("citations") or [row.get("citation_source", "")]
                )
                lines.extend([f"Figure: {caption}", figure_citations, ""])
            elif block_type == "table":
                row = table_rows.get(str(block.get("table_id", "")), {})
                source_path = row.get("source_path", "")
                columns, records = _read_table_preview(source_path)
                lines.append(f"Table: {block.get('caption') or row.get('caption', '')}")
                lines.append("| " + " | ".join(columns) + " |")
                lines.append("| " + " | ".join("---" for _ in columns) + " |")
                for record in records:
                    cells = [str(value).replace("|", "\\|") for value in record]
                    lines.append("| " + " | ".join(cells) + " |")
                lines.extend([_citations_text(block.get("citations") or [source_path]), ""])
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return output_path


def _write_docx(
    spec: dict[str, Any],
    table_rows: dict[str, dict[str, str]],
    figure_rows: dict[str, dict[str, str]],
) -> Path:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading(str(spec.get("document_title", "Model Methodology Document")), 0)
    subtitle = str(spec.get("document_subtitle", "")).strip()
    if subtitle:
        document.add_paragraph(subtitle)
    for section in sorted(spec.get("sections", []), key=lambda item: item.get("order", 9999)):
        document.add_heading(str(section.get("heading", "Untitled Section")), level=1)
        for block in section.get("blocks", []) or []:
            block_type = block.get("type")
            if block_type == "paragraph":
                paragraph = document.add_paragraph(str(block.get("text", "")).strip())
                citation_text = _citations_text(block.get("citations"))
                if citation_text:
                    paragraph.add_run(f" {citation_text}").italic = True
            elif block_type == "bullet_list":
                for item in block.get("items", []) or []:
                    document.add_paragraph(str(item), style="List Bullet")
            elif block_type == "figure":
                row = figure_rows.get(str(block.get("figure_id", "")), {})
                image_path = row.get("image_path", "")
                if image_path and (PACKAGE_ROOT / image_path).exists():
                    document.add_picture(
                        str(PACKAGE_ROOT / image_path),
                        width=Inches(float(row.get("width_inches") or 6.5)),
                    )
                caption = block.get("caption") or row.get("caption", "")
                caption_paragraph = document.add_paragraph(f"Figure. {caption}")
                citation_text = _citations_text(
                    block.get("citations") or [row.get("citation_source", "")]
                )
                if citation_text:
                    caption_paragraph.add_run(f" {citation_text}").italic = True
            elif block_type == "table":
                row = table_rows.get(str(block.get("table_id", "")), {})
                source_path = row.get("source_path", "")
                columns, records = _read_table_preview(
                    source_path,
                    max_rows=int(row.get("max_rows") or 12),
                )
                document.add_paragraph(f"Table. {block.get('caption') or row.get('caption', '')}")
                table = document.add_table(rows=1, cols=max(len(columns), 1))
                table.style = "Table Grid"
                for index, column in enumerate(columns or ["value"]):
                    table.rows[0].cells[index].text = str(column)
                for record in records:
                    cells = table.add_row().cells
                    for index, value in enumerate(record[: len(cells)]):
                        cells[index].text = str(value)[:1000]
                citation_text = _citations_text(block.get("citations") or [source_path])
                if citation_text:
                    document.add_paragraph(citation_text)
    output_path = GENERATED_ROOT / "model_methodology.docx"
    document.save(output_path)
    return output_path


def main() -> int:
    GENERATED_ROOT.mkdir(parents=True, exist_ok=True)
    if not SPEC_PATH.exists():
        raise SystemExit(
            f"Missing {SPEC_PATH}. Ask the LLM to create methodology_doc_spec.json first."
        )
    spec = _read_json(SPEC_PATH)
    table_rows = _table_manifest()
    figure_rows = _figure_manifest()
    notes: list[str] = []
    try:
        output_path = _write_docx(spec, table_rows, figure_rows)
        notes.append(f"Created DOCX: {output_path}")
    except Exception as exc:
        output_path = _write_markdown(spec, table_rows, figure_rows)
        notes.append(f"DOCX creation failed; created Markdown fallback: {output_path}")
        notes.append(f"DOCX failure detail: {exc}")
        notes.append("Install python-docx if DOCX creation is required in this environment.")
    notes_path = GENERATED_ROOT / "DOCX_BUILD_NOTES.md"
    notes_path.write_text(
        "# DOCX Build Notes\n\n" + "\n".join(f"- {note}" for note in notes) + "\n",
        encoding="utf-8",
    )
    sys.stdout.write("\n".join(notes) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
'''


def _build_docx_output_validator_script() -> str:
    return r'''"""Validate generated DOCX or Markdown output.

Usage:
    python ../08_tools/validate_docx_output.py
"""

from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path
from typing import Any


def _package_root() -> Path:
    here = Path(__file__).resolve()
    if here.parent.name == "08_tools":
        return here.parents[1]
    for parent in [here.parent, *here.parents]:
        if (parent / "EVIDENCE_INDEX.csv").exists():
            return parent
    return Path.cwd()


def validate() -> dict[str, Any]:
    package_root = _package_root()
    generated_root = package_root / "09_generated_outputs"
    docx_path = generated_root / "model_methodology.docx"
    md_path = generated_root / "model_methodology.md"
    notes_path = generated_root / "DOCX_BUILD_NOTES.md"
    errors: list[str] = []
    warnings: list[str] = []
    output_type = ""
    media_count = 0

    if docx_path.exists():
        output_type = "docx"
        try:
            with zipfile.ZipFile(docx_path) as archive:
                names = archive.namelist()
                if "word/document.xml" not in names:
                    errors.append("DOCX is missing word/document.xml.")
                media_count = len([name for name in names if name.startswith("word/media/")])
                document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
                if "{{" in document_xml or "}}" in document_xml:
                    errors.append("DOCX appears to contain unresolved template placeholders.")
        except zipfile.BadZipFile:
            errors.append("DOCX is not a valid zip-based Word file.")
    elif md_path.exists():
        output_type = "markdown_fallback"
        text = md_path.read_text(encoding="utf-8")
        if "![" in text:
            media_count = text.count("![")
        if "{{" in text or "}}" in text:
            errors.append("Markdown appears to contain unresolved template placeholders.")
        warnings.append("DOCX was not created; Markdown fallback is present.")
    else:
        errors.append("Neither model_methodology.docx nor model_methodology.md exists.")

    if not notes_path.exists():
        warnings.append("DOCX_BUILD_NOTES.md was not found.")

    result = {
        "status": "passed" if not errors else "failed",
        "output_type": output_type,
        "media_count": media_count,
        "errors": errors,
        "warnings": warnings,
    }
    (generated_root / "docx_output_validation.json").write_text(
        json.dumps(result, indent=2),
        encoding="utf-8",
    )
    return result


def main() -> int:
    result = validate()
    sys.stdout.write(json.dumps(result, indent=2) + "\n")
    return 0 if result["status"] == "passed" else 1


if __name__ == "__main__":
    raise SystemExit(main())
'''


def _build_documentation_gaps(payload: Mapping[str, Any]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    run = _mapping(payload.get("run"))
    decision = _mapping(payload.get("decision_summary"))
    if not run.get("target_column"):
        rows.append(_gap("Target definition", "missing", "Target column was not captured."))
    if not run.get("input_shape"):
        rows.append(_gap("Input population", "missing", "Input shape metadata was not captured."))
    if not decision.get("metrics"):
        rows.append(_gap("Performance metrics", "missing", "No primary metrics were captured."))
    if not payload.get("statistical_tests"):
        rows.append(
            _gap(
                "Statistical tests",
                "missing_or_empty",
                "No structured statistical-test payload was captured.",
            )
        )
    if not decision.get("validation_checklist"):
        rows.append(
            _gap(
                "Validation checklist",
                "missing",
                "No validation checklist rows were captured in the package context.",
            )
        )
    if not _feature_lineage_records(payload):
        rows.append(
            _gap(
                "Feature lineage",
                "missing_or_empty",
                "Feature lineage records were not available in package table previews.",
            )
        )
    else:
        missing_definition_count = _missing_feature_definition_count(payload)
        if missing_definition_count:
            rows.append(
                _gap(
                    "Feature dictionary",
                    "incomplete",
                    (
                        f"{missing_definition_count} feature-lineage row(s) appear to be "
                        "missing business-definition evidence."
                    ),
                )
            )
    issue_rows = [
        issue
        for issue in decision.get("issues", [])
        if isinstance(issue, Mapping) and str(issue.get("severity", "")).lower() != "none"
    ]
    for issue in issue_rows[:20]:
        rows.append(
            _gap(
                f"{issue.get('source', 'Decision issue')} - {issue.get('subject', '')}",
                str(issue.get("severity", "review")),
                str(issue.get("message", "Review the issue before relying on model output.")),
            )
        )
    for warning in list(payload.get("warnings", []))[:20]:
        rows.append(_gap("Pipeline warning", "warning", str(warning)))
    if not rows:
        rows.append(
            {
                "gap_area": "Automated package review",
                "status": "no_automated_gaps_detected",
                "description": (
                    "No package-level documentation gaps were detected. Human review is "
                    "still required."
                ),
                "recommended_action": (
                    "Validate the generated model document against source evidence and "
                    "institutional documentation standards."
                ),
            }
        )
    return rows


def _gap(area: str, status: str, description: str) -> dict[str, str]:
    return {
        "gap_area": area,
        "status": status,
        "description": description,
        "recommended_action": (
            "Do not let the LLM invent this evidence. Either add the source evidence, "
            "rerun with the needed export enabled, or disclose the limitation."
        ),
    }


def _build_documentation_gaps_markdown(rows: list[dict[str, str]]) -> str:
    lines = [
        "# Documentation Gaps",
        "",
        "Use this register to prevent the LLM from inventing missing evidence.",
        "",
    ]
    lines.extend(_markdown_table(rows))
    return "\n".join(lines).strip() + "\n"


def _build_model_type_writing_guide(payload: Mapping[str, Any]) -> str:
    run = _mapping(payload.get("run"))
    model_type = str(run.get("model_type", "")).lower()
    target_mode = str(run.get("target_mode", "")).lower()
    lines = [
        "# Model-Type-Specific Writing Guide",
        "",
        f"- Model type: `{model_type or 'n/a'}`",
        f"- Target mode: `{target_mode or 'n/a'}`",
        "",
        "## Required Writing Posture",
        "",
        "- Explain the model as configured in `config/run_config.json`.",
        "- Cite quantitative claims to package evidence.",
        "- State limitations and failed suitability checks plainly.",
        "- Do not claim approval, validation sign-off, or regulatory compliance.",
        "",
    ]
    if "scorecard" in model_type:
        lines.extend(
            [
                "## Scorecard / WoE Logistic Regression Focus",
                "",
                "- Explain binning, Weight of Evidence, Information Value, points, and overrides.",
                "- Discuss whether bins are monotonic, sparse, merged, or manually overridden.",
                "- Use feature-lineage and scorecard tables for facts; do not infer bin rationale.",
                "- Emphasize implementation transparency and reason-code usability.",
                "",
            ]
        )
    elif "logistic" in model_type or target_mode == "binary":
        lines.extend(
            [
                "## Binary Classification Focus",
                "",
                (
                    "- Explain the event definition, class balance, threshold, AUC/ROC, "
                    "KS, and calibration."
                ),
                "- Discuss events-per-feature and suitability-check evidence where present.",
                "- Separate ranking quality from calibration and threshold policy.",
                "",
            ]
        )
    elif any(token in model_type for token in ("lgd", "beta", "fractional", "tobit")):
        lines.extend(
            [
                "## LGD / Bounded-Severity Focus",
                "",
                (
                    "- Explain bounded target treatment, censoring, zero/one inflation, "
                    "or two-stage design."
                ),
                "- Discuss whether target values fit the selected model-family assumptions.",
                "- Describe severity-stage evidence separately from incidence-stage evidence.",
                "",
            ]
        )
    elif any(token in model_type for token in ("forecast", "arima", "smoothing", "components")):
        lines.extend(
            [
                "## Forecasting / Time-Series Focus",
                "",
                (
                    "- Explain the forecast horizon, temporal split design, "
                    "trend/seasonality treatment, and residual diagnostics."
                ),
                (
                    "- Discuss stationarity, structural breaks, backtesting, and "
                    "scenario context when evidence exists."
                ),
                (
                    "- Avoid saying the model is stable unless the time diagnostics "
                    "support that claim."
                ),
                "",
            ]
        )
    elif any(token in model_type for token in ("tree", "forest", "boost", "xgboost", "lightgbm")):
        lines.extend(
            [
                "## Tree / Boosting Focus",
                "",
                (
                    "- Emphasize predictive performance, feature importance, "
                    "overfitting controls, and explainability limits."
                ),
                "- Use PDP/ALE/SHAP-style evidence only when included in the package.",
                "- Avoid coefficient-style explanations unless the model provides them.",
                "",
            ]
        )
    elif any(token in model_type for token in ("survival", "cox", "aft", "hazard")):
        lines.extend(
            [
                "## Survival / Hazard Focus",
                "",
                (
                    "- Explain time-to-event framing, censoring assumptions, duration "
                    "target, and horizon interpretation."
                ),
                (
                    "- Discuss proportional-hazards or distributional assumptions only "
                    "when evidence is available."
                ),
                (
                    "- Separate hazard/risk interpretation from binary default "
                    "probability interpretation."
                ),
                "",
            ]
        )
    else:
        lines.extend(
            [
                "## General Methodology Focus",
                "",
                (
                    "- Explain model design, feature preparation, performance evidence, "
                    "limitations, and reproducibility."
                ),
                (
                    "- Use the section evidence map to decide which package files "
                    "support each section."
                ),
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def _write_prompt_variants(
    archive: ZipFile,
    added_arcnames: set[str],
    included: list[dict[str, Any]],
) -> None:
    prompts = {
        "prompt_full_methodology.md": (
            "# Prompt: Full Model Methodology\n\n"
            "Draft a complete model methodology document. Create "
            "`model_methodology.docx` if your environment supports Word file "
            "generation; otherwise create a Markdown fallback. Use the custom table of "
            "contents if supplied. Use `03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md`, "
            "`05_visual_assets/figure_placement_manifest.csv`, "
            "`02_document_planning/document_section_evidence_map.csv`, "
            "`04_evidence/claims/approved_claims.json`, "
            "`04_evidence/gaps/documentation_gaps.md`, and "
            "`06_validation_controls/citation_rules.md`. "
            "Every factual claim must cite package evidence.\n"
        ),
        "prompt_executive_summary.md": (
            "# Prompt: Executive Summary\n\n"
            "Draft a concise executive summary for model-risk stakeholders. Include "
            "purpose, model type, primary metrics, major warnings, limitations, and "
            "open review items. Cite every fact and avoid technical overreach.\n"
        ),
        "prompt_validation_challenge_review.md": (
            "# Prompt: Validation Challenge Review\n\n"
            "Act as a validation reviewer. Identify concerns, missing evidence, failed "
            "checks, unsupported conclusions, and questions that should be resolved "
            "before model reliance. Do not rewrite the model as approved.\n"
        ),
        "prompt_limitations_and_assumptions.md": (
            "# Prompt: Limitations And Assumptions\n\n"
            "Write the limitations, assumptions, and use restrictions section. Use "
            "`04_evidence/gaps/documentation_gaps.md`, decision issues, warnings, "
            "suitability checks, and validation evidence. Clearly separate known "
            "limitations from missing evidence.\n"
        ),
        "prompt_regulatory_gap_review.md": (
            "# Prompt: Regulatory Gap Review\n\n"
            "Compare the evidence package to "
            "`04_evidence/regulatory/regulatory_documentation_crosswalk.csv`. "
            "List each expected documentation area as satisfied, partial, missing, or "
            "not applicable, and cite the package evidence used for that judgment.\n"
        ),
        "prompt_1_create_document_plan.md": (
            "# Prompt 1: Create Controlled Document Plan\n\n"
            "Do not draft the methodology document yet. First, read "
            "`EVIDENCE_INDEX.csv`, `DO_NOT_CITE.md`, "
            "`00_START_HERE/MODEL_FACTS_DIGEST.md`, "
            "`02_document_planning/target_document_schema.json`, "
            "`02_document_planning/template_binding.json`, "
            "`02_document_planning/document_section_evidence_map.csv`, "
            "`04_evidence/gaps/documentation_gaps.md`, and "
            "`06_validation_controls/evidence_strength_policy.json`. "
            "Produce a section-by-section drafting "
            "plan with columns for section, objective, citable evidence, approved "
            "claims, missing evidence, human review needed, and draft status. "
            "Do not cite instruction, template, validator, or prompt files as evidence.\n"
        ),
        "prompt_2_draft_from_approved_plan.md": (
            "# Prompt 2: Draft From Approved Plan\n\n"
            "Draft the methodology document only after the document plan has been "
            "reviewed. Use `EVIDENCE_INDEX.csv` to identify citable evidence, use "
            "`04_evidence/claims/approved_claims.json` as the primary claim library, "
            "follow `02_document_planning/target_document_schema.json`, use controlled "
            "vocabulary, and cite every "
            "factual claim. Do not cite files listed in `DO_NOT_CITE.md`. "
            "Preserve all unresolved gaps. If creating DOCX output, follow "
            "`03_docx_workflow/DOCX_BUILD_INSTRUCTIONS.md` and insert only "
            "body-priority PNGs listed in "
            "`05_visual_assets/figure_placement_manifest.csv`. If DOCX is not supported, use "
            "Markdown image placeholders with captions and citations.\n"
        ),
        "prompt_3_validate_draft_against_evidence.md": (
            "# Prompt 3: Validate Draft Against Evidence\n\n"
            "Review the draft against `06_validation_controls/document_completion_rules.json`, "
            "`06_validation_controls/draft_validation_rules.json`, "
            "`06_validation_controls/document_quality_rubric.md`, "
            "`06_validation_controls/citation_coverage_validator.md`, and "
            "`06_validation_controls/unsupported_claim_detector.md`. "
            "Return citation gaps, unsupported claims, high-risk language, missing "
            "sections, and a pass/fail recommendation for human review readiness.\n"
        ),
    }
    for file_name, text in prompts.items():
        _write_text(
            archive,
            added_arcnames,
            f"{OPERATOR_INSTRUCTIONS_DIR}/prompts/{file_name}",
            text,
            included,
            evidence_area="prompt_variant",
            source="generated",
            llm_use=(
                "Task-specific prompt for LLM-assisted model-document drafting. "
                "Do not cite as model evidence."
            ),
        )


def _build_citation_rules() -> str:
    return """# Citation Rules

The LLM must follow these rules:

1. Every factual statement must cite a package source.
2. Use this citation style: `[source: package/path > field_or_section]`.
3. Use `EVIDENCE_INDEX.csv` as the authoritative citation boundary.
4. Cite only files marked `cite_as_evidence = true`.
5. Never cite files listed in `DO_NOT_CITE.md`.
6. Prefer structured sources before narrative sources when both exist.
7. Never cite raw row-level data because it is excluded by default.
8. If a source does not exist, write `Evidence not found in package`.
9. Do not cite prompts, templates, validators, rubrics, or operator instructions as evidence.
10. Do not infer approvals, owners, thresholds, or limitations that are not present.
11. Every DOCX figure caption must cite the chart path and the companion evidence
    used to interpret the chart.

Preferred source order:

1. `model_document_context.json`
2. `approved_claims.json`
3. `config/run_config.json`
4. `metadata/metrics.json`
5. `metadata/statistical_tests.json`
6. governance tables and structured table previews
7. `reports/model_development_dossier.md`
8. `reports/validation_pack.md`
9. `reports/decision_summary.md`
"""


def _write_tone_profiles(
    archive: ZipFile,
    added_arcnames: set[str],
    included: list[dict[str, Any]],
) -> None:
    profiles = {
        "tone_regulatory_formal.md": (
            "# Tone Profile: Regulatory Formal\n\n"
            "Use precise, conservative language. Prefer `the evidence indicates` over "
            "`the model proves`. Explicitly state limitations, open items, and review status.\n"
        ),
        "tone_model_development.md": (
            "# Tone Profile: Model Development\n\n"
            "Use technical but readable language for model developers. Explain design "
            "choices, data treatment, feature preparation, and performance evidence.\n"
        ),
        "tone_validation_review.md": (
            "# Tone Profile: Validation Review\n\n"
            "Use skeptical reviewer language. Separate evidence, concerns, missing "
            "information, and recommended follow-up actions.\n"
        ),
    }
    for file_name, text in profiles.items():
        _write_text(
            archive,
            added_arcnames,
            f"{OPERATOR_INSTRUCTIONS_DIR}/tone_profiles/{file_name}",
            text,
            included,
            evidence_area="tone_profile",
            source="generated",
            llm_use="Optional tone guidance for the generated model document.",
        )


def _build_human_review_checklist() -> str:
    return """# Human Review Checklist For LLM-Generated Documents

Before using an LLM-generated methodology document, a qualified reviewer should verify:

- The target definition, model purpose, and intended use match source evidence.
- The selected model type and target mode are described correctly.
- Split design and out-of-time / validation logic are not overstated.
- Every metric value and interpretation matches package evidence.
- Failed suitability checks, warnings, and documentation gaps are preserved.
- Feature definitions, transformations, imputation, binning, and lineage are factual.
- The document does not claim approval, validation sign-off, or compliance without evidence.
- The document does not reference raw data that was excluded from the package.
- The citation appendix maps major claims back to package sources.
- Institution-specific table-of-contents requirements were followed.
- DOCX figures, if present, appear in `figure_placement_manifest.csv` and have
  captions, source citations, and section relevance.
- The document follows `MODEL_DOCUMENT_STYLE_GUIDE.md` and was checked against
  `DOCX_QUALITY_CHECKLIST.md`.
"""


def _build_feature_dictionary_narrative(payload: Mapping[str, Any]) -> str:
    records = _feature_lineage_records(payload)
    lines = [
        "# Feature Dictionary Narrative",
        "",
        "Use this section as prose-ready input for feature governance and lineage writing.",
        "",
    ]
    if not records:
        lines.extend(
            [
                "No feature-lineage rows were available in the package context.",
                "The LLM should not invent source-feature definitions or transformation rationale.",
            ]
        )
        return "\n".join(lines).strip() + "\n"
    lines.extend(
        [
            f"- Feature-lineage rows available: `{len(records)}`",
            (
                "- Rows with likely missing business definitions: "
                f"`{_missing_feature_definition_count(payload)}`"
            ),
            "",
            "## Feature Lineage Preview",
            "",
        ]
    )
    for row in records[:40]:
        model_feature = _first_present(row, ["model_feature_name", "feature_name", "feature"])
        source_feature = _first_present(row, ["source_feature", "raw_feature", "source"])
        lineage_type = _first_present(row, ["lineage_type", "transformation_type", "type"])
        definition = _first_present(
            row,
            ["business_definition", "definition", "description", "documentation"],
        )
        guidance = _first_present(row, ["review_guidance", "selection_rationale", "rationale"])
        lines.append(f"- `{model_feature or 'n/a'}`")
        lines.append(f"  Source feature: `{source_feature or 'n/a'}`.")
        lines.append(f"  Lineage type: `{lineage_type or 'n/a'}`.")
        if definition:
            lines.append(f"  Business definition: {definition}")
        else:
            lines.append("  Business definition: Evidence not found in package.")
        if guidance:
            lines.append(f"  Review guidance: {guidance}")
    return "\n".join(lines).strip() + "\n"


def _build_metrics_interpretation_brief(payload: Mapping[str, Any]) -> str:
    decision = _mapping(payload.get("decision_summary"))
    metric_rows = [row for row in decision.get("metrics", []) if isinstance(row, Mapping)]
    lines = [
        "# Metrics Interpretation Brief",
        "",
        "This brief gives the LLM prose-ready metric language with source constraints.",
        "",
    ]
    if not metric_rows:
        lines.append("No metric rows were captured. The LLM must state that evidence is missing.")
        return "\n".join(lines).strip() + "\n"
    for row in metric_rows:
        split = row.get("split", "n/a")
        metric = row.get("metric", "n/a")
        value = row.get("value", "n/a")
        interpretation = row.get("interpretation", "contextual")
        lines.append(
            f"- `{metric}` on `{split}` was `{value}` and is interpreted as "
            f"`{interpretation}`. "
            "[source: model_document_context.json > decision_summary.metrics]"
        )
    lines.extend(
        [
            "",
            "Writing rules:",
            "",
            "- Do not compare metrics across runs unless comparison evidence is included.",
            "- Do not describe calibration as strong unless calibration evidence is present.",
            "- Use `watch` or `bad` interpretations as open review items, not approvals.",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _build_figure_placement_manifest(included: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, dict[str, str]] = {}
    for record in included:
        package_path = str(record.get("package_path", ""))
        if (
            record.get("source") != "generated_chart_export"
            or not _is_visual_asset_path(package_path)
        ):
            continue
        path = Path(package_path)
        suffix = path.suffix.lower().lstrip(".")
        if suffix not in {"png", "html"}:
            continue
        stem = path.stem
        grouped.setdefault(stem, {})[suffix] = package_path

    if not grouped:
        return [
            {
                "figure_rank": "",
                "figure_id": "",
                "figure_name": "no_generated_chart_assets",
                "image_path": "",
                "html_review_path": "",
                "file_format": "",
                "recommended_section": "Appendices And Source Evidence Index",
                "caption": "No generated chart assets were included in this package.",
                "why_include": "No chart insertion is available for the DOCX draft.",
                "source_chart": "",
                "citation_source": "",
                "priority": "",
                "width_inches": "",
                "insert_or_appendix": "not_available",
                "docx_instruction": "Do not invent figures. Draft from available tables and text.",
            }
        ]

    rows: list[dict[str, Any]] = []
    ordered_names = sorted(grouped, key=lambda name: (_document_figure_priority(name), name))
    for position, figure_name in enumerate(ordered_names, start=1):
        assets = grouped[figure_name]
        metadata = _document_figure_metadata(figure_name)
        preferred_path = assets.get("png") or assets.get("html", "")
        preferred_format = "png" if assets.get("png") else "html"
        body_priority = bool(assets.get("png")) and position <= 8
        figure_id = _document_asset_id("FIG", position, figure_name)
        rows.append(
            {
                "figure_rank": position,
                "figure_id": figure_id,
                "figure_name": figure_name,
                "image_path": preferred_path,
                "html_review_path": assets.get("html", ""),
                "file_format": preferred_format,
                "recommended_section": metadata["recommended_section"],
                "caption": metadata["caption"],
                "why_include": metadata["why_include"],
                "source_chart": preferred_path,
                "citation_source": preferred_path,
                "priority": metadata["priority"],
                "width_inches": "6.5" if body_priority else "7.0",
                "insert_or_appendix": "body" if body_priority else "appendix",
                "docx_instruction": (
                    "Insert the PNG in the main body with a figure number and caption."
                    if body_priority
                    else (
                        "Use as appendix support or omit unless it directly supports a "
                        "document claim."
                    )
                ),
            }
        )
    return rows


def _is_visual_asset_path(package_path: str) -> bool:
    path_lower = package_path.lower().replace("\\", "/")
    figure_root = FIGURE_ASSETS_ROOT.lower().replace("\\", "/")
    return (
        f"/{figure_root}/" in f"/{path_lower}"
        or "/source_artifacts/figures/" in path_lower
    )


def _document_asset_id(prefix: str, position: int, name: str) -> str:
    safe = _safe_name(name).replace(".", "_").replace("-", "_").upper()
    safe = "_".join(part for part in safe.split("_") if part)
    return f"{prefix}_{position:03d}_{safe[:48] or 'ASSET'}"


def _build_figure_contact_sheet(rows: list[dict[str, Any]]) -> str:
    lines = [
        "# Figure Contact Sheet",
        "",
        "Use this index to decide which visuals belong in the methodology document.",
        "The manifest CSV remains the authority for exact IDs and paths.",
        "",
    ]
    available_rows = [row for row in rows if row.get("figure_id")]
    if not available_rows:
        lines.append("No document-ready chart images were included in this package.")
        return "\n".join(lines).strip() + "\n"
    for row in available_rows:
        figure_id = row.get("figure_id", "")
        title = str(row.get("figure_name", "")).replace("_", " ").title()
        image_path = str(row.get("image_path", ""))
        relative_image = _relative_visual_path_for_contact_sheet(image_path)
        lines.extend(
            [
                f"## {figure_id}: {title}",
                "",
                f"- Recommended section: {row.get('recommended_section', '')}",
                f"- Placement: {row.get('insert_or_appendix', '')}",
                f"- Caption: {row.get('caption', '')}",
                f"- Citation source: `{row.get('citation_source', '')}`",
                "",
            ]
        )
        if relative_image and image_path.lower().endswith((".png", ".jpg", ".jpeg", ".svg")):
            lines.extend([f"![{title}]({relative_image})", ""])
        else:
            lines.extend(
                [
                    "No PNG image is available for inline preview. Use the HTML review path:",
                    f"`{row.get('html_review_path', '')}`",
                    "",
                ]
            )
    return "\n".join(lines).strip() + "\n"


def _relative_visual_path_for_contact_sheet(package_path: str) -> str:
    normalized = package_path.replace("\\", "/")
    prefix = f"{VISUAL_ASSETS_DIR}/"
    if normalized.startswith(prefix):
        return normalized[len(prefix) :]
    return normalized


def _build_table_placement_manifest(included: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates = [
        record
        for record in included
        if _is_document_table_candidate(record)
    ]
    ordered_candidates = sorted(
        candidates,
        key=lambda record: (
            _document_table_priority(str(record.get("package_path", ""))),
            str(record.get("package_path", "")),
        ),
    )
    if not ordered_candidates:
        return [
            {
                "table_rank": "",
                "table_id": "",
                "table_name": "no_table_assets",
                "source_path": "",
                "recommended_section": "Appendices And Source Evidence Index",
                "title": "No high-value table assets were identified.",
                "caption": "No table insertion is available for the DOCX draft.",
                "citation_source": "",
                "priority": "",
                "max_rows": "",
                "insert_or_appendix": "not_available",
                "docx_instruction": (
                    "Do not invent tables. Draft from available text and citable evidence."
                ),
            }
        ]

    rows: list[dict[str, Any]] = []
    for position, record in enumerate(ordered_candidates[:40], start=1):
        package_path = str(record.get("package_path", ""))
        table_name = Path(package_path).stem
        metadata = _document_table_metadata(package_path)
        table_id = _document_asset_id("TABLE", position, table_name)
        rows.append(
            {
                "table_rank": position,
                "table_id": table_id,
                "table_name": table_name,
                "source_path": package_path,
                "recommended_section": metadata["recommended_section"],
                "title": metadata["title"],
                "caption": metadata["caption"],
                "citation_source": package_path,
                "priority": metadata["priority"],
                "max_rows": metadata["max_rows"],
                "insert_or_appendix": metadata["insert_or_appendix"],
                "docx_instruction": metadata["docx_instruction"],
            }
        )
    return rows


def _is_document_table_candidate(record: Mapping[str, Any]) -> bool:
    package_path = str(record.get("package_path", ""))
    path_lower = package_path.lower().replace("\\", "/")
    suffix = Path(package_path).suffix.lower()
    if suffix not in {".csv", ".json"}:
        return False
    if any(
        token in path_lower
        for token in (
            "/07_operator_prompts/",
            "/08_tools/",
            "/01_document_template/",
            "schema.json",
            "policy.json",
            "rules.json",
            "rubric.json",
            "manifest.json",
            "plotly.min.js",
            "figure_manifest.json",
        )
    ):
        return False
    evidence_area = str(record.get("evidence_area", "")).lower()
    if evidence_area in {
        "llm_prompt",
        "docx_authoring_control",
        "draft_validation",
        "document_quality",
        "document_control",
        "citation_rules",
        "document_template",
        "manifest",
    }:
        return False
    if "source_citation_map.csv" in path_lower or "evidence_index.csv" in path_lower:
        return False
    return any(
        token in path_lower
        for token in (
            "metrics",
            "statistical_tests",
            "validation_checklist",
            "traceability",
            "feature_lineage",
            "feature_importance",
            "documentation_gaps",
            "approved_claims",
            "regulatory_documentation_crosswalk",
            "document_section_evidence_map",
            "model_document_context",
            "table_previews",
            "converted_tables",
            "/tables/",
            "governance",
        )
    )


def _document_table_priority(package_path: str) -> int:
    normalized = package_path.lower().replace("-", "_").replace(" ", "_")
    priorities = (
        ("metric", "performance"),
        ("validation_checklist",),
        ("statistical_tests",),
        ("feature_importance", "feature_lineage"),
        ("documentation_gaps",),
        ("approved_claims",),
        ("regulatory_documentation_crosswalk",),
        ("document_section_evidence_map",),
        ("traceability",),
        ("model_document_context",),
    )
    for priority, patterns in enumerate(priorities, start=1):
        if any(pattern in normalized for pattern in patterns):
            return priority
    return 99


def _document_table_metadata(package_path: str) -> dict[str, Any]:
    normalized = package_path.lower().replace("-", "_").replace(" ", "_")
    table_title = Path(package_path).stem.replace("_", " ").title()
    if "metric" in normalized or "performance" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "title": table_title,
            "caption": "Primary model-performance evidence for the methodology document.",
            "priority": 1,
            "max_rows": 25,
            "insert_or_appendix": "body",
            "docx_instruction": "Insert as a concise body table with rounded metrics.",
        }
    if "validation_checklist" in normalized:
        return {
            "recommended_section": "Governance, Reproducibility, And Controls",
            "title": table_title,
            "caption": "Validation checklist evidence that should be preserved for review.",
            "priority": 2,
            "max_rows": 30,
            "insert_or_appendix": "body",
            "docx_instruction": (
                "Insert failed or warning rows in the body; move full table to appendix."
            ),
        }
    if "statistical_tests" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "title": table_title,
            "caption": "Statistical-test evidence for model validation discussion.",
            "priority": 3,
            "max_rows": 30,
            "insert_or_appendix": "body",
            "docx_instruction": (
                "Summarize material tests in the body and cite the full source table."
            ),
        }
    if "feature" in normalized:
        return {
            "recommended_section": "Features, Transformations, Imputation, And Lineage",
            "title": table_title,
            "caption": "Feature lineage, importance, or transformation evidence.",
            "priority": 4,
            "max_rows": 30,
            "insert_or_appendix": "body",
            "docx_instruction": (
                "Insert the most material feature rows; appendix any full inventory."
            ),
        }
    if "documentation_gaps" in normalized:
        return {
            "recommended_section": "Assumptions, Limitations, And Open Review Items",
            "title": table_title,
            "caption": "Open documentation gaps and limitations that must not be hidden.",
            "priority": 5,
            "max_rows": 40,
            "insert_or_appendix": "body",
            "docx_instruction": "Insert all material gaps in the body limitations section.",
        }
    if "approved_claims" in normalized:
        return {
            "recommended_section": "Executive Summary",
            "title": table_title,
            "caption": "Evidence-backed claim library for controlled drafting.",
            "priority": 6,
            "max_rows": 20,
            "insert_or_appendix": "appendix",
            "docx_instruction": (
                "Use as drafting evidence; do not paste the full claim library in body."
            ),
        }
    if "regulatory" in normalized or "crosswalk" in normalized:
        return {
            "recommended_section": "Governance, Reproducibility, And Controls",
            "title": table_title,
            "caption": "Regulatory documentation crosswalk for review traceability.",
            "priority": 7,
            "max_rows": 30,
            "insert_or_appendix": "appendix",
            "docx_instruction": (
                "Use as appendix support unless a governance section needs the table."
            ),
        }
    return {
        "recommended_section": "Appendices And Source Evidence Index",
        "title": table_title,
        "caption": "Supporting table evidence from the completed run.",
        "priority": 99,
        "max_rows": 25,
        "insert_or_appendix": "appendix",
        "docx_instruction": "Use in appendix or cite only if it supports a specific claim.",
    }


def _document_figure_priority(name: str) -> int:
    normalized = name.lower().replace("-", "_").replace(" ", "_")
    for priority, patterns in enumerate(DOCUMENT_FIGURE_PRIORITY_PATTERNS):
        if any(pattern in normalized for pattern in patterns):
            return priority
    return len(DOCUMENT_FIGURE_PRIORITY_PATTERNS)


def _document_figure_metadata(name: str) -> dict[str, str | int]:
    normalized = name.lower().replace("-", "_").replace(" ", "_")
    title = name.replace("_", " ").title()
    if "roc" in normalized or "auc" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} summarizes ranking performance. Use it with metric tables to "
                "support discrimination discussion."
            ),
            "why_include": "High-value performance evidence for binary model discrimination.",
            "priority": 1,
        }
    if "ks" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} summarizes event/non-event separation. Use it with KS table "
                "evidence and threshold discussion."
            ),
            "why_include": "Common credit-risk validation evidence for score separation.",
            "priority": 2,
        }
    if "calibration" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} compares predicted risk to observed outcomes. Use it to discuss "
                "probability alignment and calibration limitations."
            ),
            "why_include": "Key evidence for probability model reliability.",
            "priority": 3,
        }
    if "threshold" in normalized or "lift" in normalized or "gain" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} shows decision-threshold or ranking tradeoffs. Use it only with "
                "the selected threshold policy and companion metrics."
            ),
            "why_include": "Useful for decision cutoff, lift, gain, and operating-point review.",
            "priority": 4,
        }
    if any(token in normalized for token in ("feature_importance", "driver", "importance")):
        return {
            "recommended_section": "Features, Transformations, Imputation, And Lineage",
            "caption": (
                f"{title} summarizes the strongest model drivers. Use it with feature "
                "lineage evidence and avoid unsupported causal interpretations."
            ),
            "why_include": "Supports feature-driver and explainability discussion.",
            "priority": 5,
        }
    if "partial_dependence" in normalized or "pdp" in normalized:
        return {
            "recommended_section": "Assumptions, Limitations, And Open Review Items",
            "caption": (
                f"{title} shows modeled response sensitivity. Use it as directional "
                "explainability evidence, not as causal proof."
            ),
            "why_include": "Helpful for explainability and sensitivity discussion.",
            "priority": 6,
        }
    if any(token in normalized for token in ("scorecard", "woe", "bin")):
        return {
            "recommended_section": "Features, Transformations, Imputation, And Lineage",
            "caption": (
                f"{title} summarizes binning or scorecard behavior. Use it with WoE/IV "
                "tables and scorecard lineage."
            ),
            "why_include": "Useful for scorecard transparency and binning review.",
            "priority": 7,
        }
    if "actual_vs_predicted" in normalized or "residual" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} summarizes prediction error patterns. Use it with residual "
                "diagnostics and limitations."
            ),
            "why_include": "Useful for continuous model fit and residual review.",
            "priority": 8,
        }
    if "psi" in normalized or "stability" in normalized:
        return {
            "recommended_section": "Performance, Calibration, Stability, And Statistical Testing",
            "caption": (
                f"{title} summarizes stability evidence. Use it to discuss population or "
                "feature shift where supported by companion tables."
            ),
            "why_include": "Supports stability and drift discussion.",
            "priority": 9,
        }
    return {
        "recommended_section": "Appendices And Source Evidence Index",
        "caption": f"{title} is supporting visual evidence from the completed run.",
        "why_include": "Supporting visual evidence; include only if tied to a specific claim.",
        "priority": 99,
    }


def _build_chart_interpretation_brief(
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
) -> str:
    chart_rows = [
        row
        for row in included
        if row.get("evidence_area") == "visual_evidence"
        or "/figures/" in str(row.get("package_path", ""))
    ]
    skipped_charts = [
        row
        for row in skipped
        if str(row.get("artifact_key", "")) == "figures_dir"
        or "chart" in str(row.get("reason", "")).lower()
    ]
    lines = [
        "# Chart Interpretation Brief",
        "",
        "Use this brief to discuss charts without hallucinating visual conclusions.",
        "Use `figure_placement_manifest.csv` to decide which PNG files should be "
        "inserted into a DOCX methodology document.",
        "",
    ]
    if chart_rows:
        lines.extend(["## Included Chart Files", ""])
        for row in chart_rows[:80]:
            lines.append(
                f"- `{row.get('package_path', '')}`: {row.get('llm_use', 'Visual evidence.')}"
            )
    else:
        lines.append(
            "No standalone chart files were included. Use narrative chart descriptions, "
            "diagnostic tables, and the interactive report reference if present."
        )
    if skipped_charts:
        lines.extend(["", "## Skipped Chart Files", ""])
        for row in skipped_charts[:40]:
            lines.append(
                f"- `{row.get('original_path', '')}` skipped because "
                f"{row.get('reason', '')}."
            )
    lines.extend(
        [
            "",
            "Interpretation rules:",
            "",
            "- Cite chart files or chart-related tables when making visual claims.",
            "- Do not infer trend, calibration, or stability conclusions from file names alone.",
            (
                "- If chart evidence is missing, state that the chart was not included "
                "in the LLM package."
            ),
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _source_status(payload: Mapping[str, Any], sources: list[str]) -> str:
    statuses = [_source_available(payload, source) for source in sources]
    if all(statuses):
        return "present"
    if any(statuses):
        return "partial"
    return "not_found"


def _source_available(payload: Mapping[str, Any], source: str) -> bool:
    source_lower = source.lower().replace("*", "")
    generated_tokens = (
        "model_document_context",
        "approved_claims",
        "documentation_gaps",
        "document_section_evidence_map",
        "regulatory_documentation_crosswalk",
        "model_type_writing_guide",
        "metrics_interpretation_brief",
        "chart_interpretation_brief",
        "citation_rules",
        "controlled_vocabulary",
        "document_completion_rules",
        "document_quality_rubric",
        "human_review_checklist",
        "evidence_strength_policy",
        "draft_validation_rules",
        "llm_evidence_checklist",
        "llm_redaction_policy",
        "source_citation_map",
        "target_document_schema",
        "template_binding",
        "default_model_methodology_outline",
        "regulatory_language_guardrails",
        "citation_coverage_validator",
        "unsupported_claim_detector",
    )
    if any(token in source_lower for token in generated_tokens):
        return True
    if "structured_table_previews" in source_lower or "table_previews" in source_lower:
        return bool(_mapping(payload.get("diagnostic_table_previews")))
    artifact_paths = " ".join(
        str(value).lower().replace("\\", "/")
        for value in _mapping(payload.get("artifacts")).values()
    )
    source_name = Path(source_lower).name
    if source_lower and source_lower in artifact_paths:
        return True
    if source_name and source_name in artifact_paths:
        return True
    diagnostic_tables = _mapping(payload.get("diagnostic_table_previews"))
    return any(str(table_name).lower() in source_lower for table_name in diagnostic_tables)


def _feature_lineage_records(payload: Mapping[str, Any]) -> list[Mapping[str, Any]]:
    table_previews = _mapping(payload.get("diagnostic_table_previews"))
    for table_name, table_payload in table_previews.items():
        if "feature_lineage" in str(table_name).lower():
            return [
                row
                for row in _mapping(table_payload).get("records", [])
                if isinstance(row, Mapping)
            ]
    return []


def _missing_feature_definition_count(payload: Mapping[str, Any]) -> int:
    missing_count = 0
    for row in _feature_lineage_records(payload):
        definition = _first_present(
            row,
            ["business_definition", "definition", "description", "documentation"],
        )
        if not definition or str(definition).strip().lower() in {"", "nan", "none", "n/a"}:
            missing_count += 1
    return missing_count


def _first_present(row: Mapping[str, Any], candidates: list[str]) -> str:
    lower_lookup = {str(key).lower(): value for key, value in row.items()}
    for candidate in candidates:
        value = lower_lookup.get(candidate.lower())
        if value is None:
            continue
        text = str(value).strip()
        if text and text.lower() not in {"nan", "none", "n/a"}:
            return text
    return ""


def _build_guidance_mapping() -> list[dict[str, Any]]:
    return [
        {
            "guidance_area": "Model development, implementation, and use",
            "documentation_need": (
                "Purpose, intended use, model design, data, assumptions, and limitations."
            ),
            "primary_package_sources": [
                "model_document_context.md",
                "reports/model_development_dossier.md",
                "reports/model_documentation_pack.md",
                "config/run_config.json",
                "code/generated_run.py",
            ],
        },
        {
            "guidance_area": "Model validation",
            "documentation_need": (
                "Performance, conceptual soundness evidence, process verification, "
                "outcomes analysis, and limitations."
            ),
            "primary_package_sources": [
                "reports/validation_pack.md",
                "reports/decision_summary.md",
                "metadata/metrics.json",
                "metadata/statistical_tests.json",
                "tables/governance/validation_checklist.*",
                "tables/governance/evidence_traceability_map.*",
            ],
        },
        {
            "guidance_area": "Governance, policies, and controls",
            "documentation_need": (
                "Reproducibility, run controls, source evidence, lineage, "
                "configuration, and audit trail."
            ),
            "primary_package_sources": [
                "artifact_manifest.json",
                "metadata/reproducibility_manifest.json",
                "metadata/step_manifest.json",
                "metadata/run_debug_trace.json",
                "tables/governance/feature_lineage_map.*",
            ],
        },
        {
            "guidance_area": "CECL / CCAR supporting context when applicable",
            "documentation_need": (
                "Forecasting, segmentation, scenario, horizon, and loss-definition "
                "context where relevant to the configured model."
            ),
            "primary_package_sources": [
                "config/run_config.json",
                "model_document_context.json",
                "reports/model_development_dossier.md",
                "04_evidence/table_previews/",
            ],
        },
    ]


def _build_evidence_checklist(
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
) -> list[dict[str, str]]:
    included_text = "\n".join(
        (
            f"{row.get('artifact_key', '')} {row.get('package_path', '')} "
            f"{row.get('original_path', '')}"
        )
        for row in included
    ).lower()
    skipped_text = "\n".join(
        f"{row.get('artifact_key', '')} {row.get('original_path', '')} {row.get('reason', '')}"
        for row in skipped
    ).lower()
    requirements = [
        ("Run configuration", "run_config.json", "Required to document selected settings."),
        ("Generated Python workflow", "generated_run.py", "Required for reproducibility."),
        ("Decision summary", "decision_summary.md", "Required for recommendation context."),
        (
            "Model development dossier",
            "model_development_dossier.md",
            "Required for methodology narrative.",
        ),
        ("Validation pack", "validation_pack.md", "Required for validation evidence."),
        ("Metrics", "metrics.json", "Required for quantitative performance evidence."),
        ("Statistical tests", "statistical_tests.json", "Required for test evidence."),
        (
            "Feature lineage",
            "feature_lineage_map",
            "Required for source-to-model feature traceability.",
        ),
        (
            "Reproducibility manifest",
            "reproducibility_manifest.json",
            "Required for environment and hash evidence.",
        ),
        ("Artifact manifest", "artifact_manifest.json", "Required for source file indexing."),
        (
            "Section evidence map",
            "document_section_evidence_map",
            "Required to map document sections to package evidence.",
        ),
        (
            "Approved claims",
            "approved_claims",
            "Required to reduce unsupported LLM-generated claims.",
        ),
        (
            "Documentation gaps",
            "documentation_gaps",
            "Required to prevent the LLM from inventing missing evidence.",
        ),
        (
            "Regulatory crosswalk",
            "regulatory_documentation_crosswalk",
            "Useful for SR 11-7/OCC-style methodology document structure.",
        ),
        (
            "Model-specific writing guide",
            "model_type_writing_guide",
            "Useful for model-family-specific methodology language.",
        ),
        (
            "Citation rules",
            "citation_rules",
            "Required to force source-backed factual statements.",
        ),
        (
            "Target document schema",
            "target_document_schema",
            "Required to control expected document sections and evidence use.",
        ),
        (
            "Evidence strength policy",
            "evidence_strength_policy",
            "Required to rank sources and prevent weak evidence from supporting strong claims.",
        ),
        (
            "Document completion rules",
            "document_completion_rules",
            "Required to define blocking failures before human review.",
        ),
        (
            "Controlled vocabulary",
            "controlled_vocabulary",
            "Required to prevent overstatement and standardize model-risk language.",
        ),
        (
            "Draft validation rules",
            "draft_validation_rules",
            "Required to check LLM output for missing citations and unsupported claims.",
        ),
        (
            "Document quality rubric",
            "document_quality_rubric",
            "Useful for consistent review scoring of LLM-generated drafts.",
        ),
        (
            "Template binding",
            "template_binding",
            "Useful for mapping institution-specific templates to package evidence.",
        ),
        (
            "Redaction policy",
            "llm_redaction_policy",
            "Required to keep LLM use aligned with minimum necessary evidence.",
        ),
        (
            "Metrics interpretation brief",
            "metrics_interpretation_brief",
            "Useful for performance section drafting.",
        ),
        (
            "Chart interpretation brief",
            "chart_interpretation_brief",
            "Useful for visual evidence discussion.",
        ),
        (
            "Figure placement manifest",
            "figure_placement_manifest",
            "Required when generating DOCX output with inserted charts.",
        ),
        (
            "DOCX build instructions",
            "docx_build_instructions",
            "Required when asking an LLM to generate Word-ready output.",
        ),
        (
            "Human review checklist",
            "human_review_checklist",
            "Required before using an LLM-generated document.",
        ),
        (
            "Diagnostic table previews",
            "table_previews",
            "Useful for LLM-readable diagnostics.",
        ),
        ("Chart exports", "figures/", "Useful when individual chart exports exist."),
        (
            "Table of contents placeholder",
            "drop_table_of_contents_here",
            "Supports custom document templates.",
        ),
    ]
    rows: list[dict[str, str]] = []
    for area, token, rationale in requirements:
        token_lower = token.lower()
        present = token_lower in included_text
        rows.append(
            {
                "evidence_area": area,
                "status": "present" if present else "not_found",
                "why_it_matters": rationale,
                "recommended_action": (
                    "Use the included evidence."
                    if present
                    else (
                        "Review skipped files and run artifacts; regenerate the model "
                        "run if needed."
                    )
                ),
            }
        )
    rows.append(
        {
            "evidence_area": "Privacy exclusions",
            "status": "applied" if "row-level" in skipped_text else "not_triggered",
            "why_it_matters": (
                "Raw data and row-level predictions should not be sent to an LLM by default."
            ),
            "recommended_action": (
                "Only add row-level data under an approved privacy and governance process."
            ),
        }
    )
    return rows


def _build_evidence_index(
    included: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for record in included:
        classification = _classify_included_record(record)
        rows.append(
            {
                "package_path": record.get("package_path", ""),
                "status": "included",
                "category": classification["category"],
                "cite_as_evidence": classification["cite_as_evidence"],
                "citation_boundary": classification["citation_boundary"],
                "evidence_area": record.get("evidence_area", ""),
                "source": record.get("source", ""),
                "artifact_key": record.get("artifact_key", ""),
                "size_bytes": record.get("size_bytes", 0),
                "llm_use": record.get("llm_use", ""),
                "do_not_cite_reason": classification["do_not_cite_reason"],
            }
        )
    for record in skipped:
        rows.append(
            {
                "package_path": "",
                "status": "skipped",
                "category": "do_not_cite",
                "cite_as_evidence": False,
                "citation_boundary": "not_available",
                "evidence_area": "",
                "source": "",
                "artifact_key": record.get("artifact_key", ""),
                "size_bytes": record.get("size_bytes", 0),
                "llm_use": "",
                "original_path": record.get("original_path", ""),
                "do_not_cite_reason": record.get("reason", "Skipped file is not available."),
            }
        )
    return rows


def _build_do_not_cite_markdown(evidence_index_rows: list[dict[str, Any]]) -> str:
    do_not_cite_rows = [
        row
        for row in evidence_index_rows
        if not _truthy(row.get("cite_as_evidence")) and row.get("status") == "included"
    ]
    lines = [
        "# Do Not Cite As Model Evidence",
        "",
        (
            "The files below are useful workflow controls, prompts, templates, "
            "validators, or support files. They should guide LLM behavior, but they "
            "must not be cited as factual evidence about the model run."
        ),
        "",
        "Use `EVIDENCE_INDEX.csv` and `source_citation_map.csv` to identify citable evidence.",
        "",
    ]
    if not do_not_cite_rows:
        lines.append("No included files were classified as non-citable.")
        return "\n".join(lines).strip() + "\n"
    lines.extend(["| Package Path | Category | Reason |", "| --- | --- | --- |"])
    for row in do_not_cite_rows:
        lines.append(
            "| "
            f"`{row.get('package_path', '')}` | "
            f"{row.get('category', '')} | "
            f"{row.get('do_not_cite_reason', '')} |"
        )
    return "\n".join(lines).strip() + "\n"


def _citable_records(included: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        record
        for record in included
        if _classify_included_record(record)["cite_as_evidence"]
    ]


def _classify_included_record(record: Mapping[str, Any]) -> dict[str, Any]:
    package_path = str(record.get("package_path", ""))
    path_lower = package_path.lower()
    evidence_area = str(record.get("evidence_area", "")).lower()
    source = str(record.get("source", "")).lower()

    category = "evidence"
    cite_as_evidence = True
    reason = ""

    if _is_visual_asset_path(package_path):
        if path_lower.endswith(("plotly.min.js", "figure_manifest.json")):
            category = "do_not_cite"
            cite_as_evidence = False
            reason = "Chart support file or manifest, not model evidence."
        else:
            category = "supporting_chart"
            cite_as_evidence = True
    elif (
        "/07_operator_prompts/" in path_lower
        or "/operator_instructions/" in path_lower
        or "prompt" in path_lower
    ):
        category = "instruction"
        cite_as_evidence = False
        reason = "Operator prompt/instruction file, not model-run evidence."
    elif (
        "/01_document_template/" in path_lower
        or "/document_template/" in path_lower
        or path_lower.endswith(
        (
            "default_model_methodology_outline.md",
            "target_document_schema.json",
            "template_binding.json",
            "resolved_toc_binding.schema.json",
            "methodology_doc_spec.schema.json",
            "methodology_doc_spec_example.json",
        )
        )
    ):
        category = "template"
        cite_as_evidence = False
        reason = "Document structure or template control, not model-run evidence."
    elif "/08_tools/" in path_lower or "/tools/" in path_lower or any(
        token in path_lower
        for token in (
            "validator",
            "unsupported_claim_detector",
            "document_quality_rubric",
            "draft_validation_rules",
        )
    ):
        category = "validator"
        cite_as_evidence = False
        reason = "Draft-validation aid, not model-run evidence."
    elif any(
        token in path_lower
        for token in (
            "readme_llm_package",
            "llm_evidence_manifest",
            "llm_package_build_profile",
            "evidence_index",
            "do_not_cite",
            "docx_build_instructions",
            "docx_quality_checklist",
            "model_document_style_guide",
            "figure_placement_manifest",
            "figure_contact_sheet",
            "table_placement_manifest",
            "resolved_toc_binding_instructions",
            "citation_rules",
            "evidence_strength_policy",
            "document_completion_rules",
            "controlled_vocabulary",
            "regulatory_language_guardrails",
            "llm_redaction_policy",
            "human_review_checklist",
        )
    ):
        category = "instruction"
        cite_as_evidence = False
        reason = "Package control or use-policy file, not factual model evidence."
    elif evidence_area in {"orientation", "citation_map", "evidence_checklist"}:
        category = "do_not_cite"
        cite_as_evidence = False
        reason = "Package navigation/control file, not underlying model evidence."
    elif source in {"generated_chart_export"} and path_lower.endswith(".json"):
        category = "do_not_cite"
        cite_as_evidence = False
        reason = "Generated chart manifest, not model evidence."

    return {
        "category": category,
        "cite_as_evidence": bool(cite_as_evidence),
        "citation_boundary": "citable" if cite_as_evidence else "do_not_cite",
        "do_not_cite_reason": reason,
    }


def _truthy(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"true", "1", "yes", "y"}


def _build_checklist_markdown(rows: list[dict[str, str]]) -> str:
    lines = ["# LLM Evidence Checklist", ""]
    lines.extend(_markdown_table(rows))
    return "\n".join(lines).strip() + "\n"


def _write_text(
    archive: ZipFile,
    added_arcnames: set[str],
    arcname: str,
    text: str,
    included: list[dict[str, Any]],
    *,
    evidence_area: str,
    source: str,
    llm_use: str = "Generated package guidance or model-document context.",
) -> None:
    if arcname in added_arcnames:
        return
    archive.writestr(arcname, text)
    added_arcnames.add(arcname)
    included.append(
        {
            "package_path": arcname,
            "original_path": "",
            "artifact_key": "",
            "evidence_area": evidence_area,
            "source": source,
            "size_bytes": len(text.encode("utf-8")),
            "llm_use": llm_use,
        }
    )


def _write_json(
    archive: ZipFile,
    added_arcnames: set[str],
    arcname: str,
    payload: Any,
    included: list[dict[str, Any]],
    *,
    evidence_area: str,
    source: str,
) -> None:
    _write_text(
        archive,
        added_arcnames,
        arcname,
        json.dumps(payload, indent=2, default=str),
        included,
        evidence_area=evidence_area,
        source=source,
        llm_use="Structured JSON context for LLM ingestion.",
    )


def _write_csv(
    archive: ZipFile,
    added_arcnames: set[str],
    arcname: str,
    rows: list[dict[str, Any]],
    included: list[dict[str, Any]],
    *,
    evidence_area: str,
    source: str,
    llm_use: str = "CSV evidence table for LLM ingestion and citation.",
) -> None:
    if arcname in added_arcnames:
        return
    text = _rows_to_csv(rows)
    archive.writestr(arcname, text)
    added_arcnames.add(arcname)
    included.append(
        {
            "package_path": arcname,
            "original_path": "",
            "artifact_key": "",
            "evidence_area": evidence_area,
            "source": source,
            "size_bytes": len(text.encode("utf-8")),
            "llm_use": llm_use,
        }
    )


def _rows_to_csv(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    fieldnames = sorted({str(key) for row in rows for key in row})
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames, lineterminator="\n")
    writer.writeheader()
    for row in rows:
        writer.writerow({key: _csv_value(row.get(key, "")) for key in fieldnames})
    return output.getvalue()


def _csv_value(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, list | tuple | set | dict):
        return json.dumps(value, default=str)
    return value


def _markdown_table(rows: Any) -> list[str]:
    if not rows:
        return ["No records available."]
    records = list(rows)
    columns = sorted({str(key) for row in records if isinstance(row, Mapping) for key in row})
    if not columns:
        return ["No records available."]
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in records[:50]:
        if not isinstance(row, Mapping):
            continue
        lines.append(
            "| "
            + " | ".join(_markdown_cell(row.get(column, "")) for column in columns)
            + " |"
        )
    if len(records) > 50:
        lines.append(f"\nShowing 50 of {len(records):,} records.")
    return lines


def _markdown_cell(value: Any) -> str:
    return str(_csv_value(value)).replace("\n", " ").replace("|", "\\|")


def _frame_payload(value: Any) -> dict[str, Any]:
    if not isinstance(value, pd.DataFrame):
        return {"rows": 0, "columns": [], "records": []}
    frame = value.iloc[
        :DEFAULT_MAX_TABLE_PREVIEW_ROWS,
        :DEFAULT_MAX_TABLE_PREVIEW_COLUMNS,
    ].copy()
    return {
        "rows": int(value.shape[0]),
        "columns": [str(column) for column in value.columns],
        "preview_rows": int(frame.shape[0]),
        "preview_columns": [str(column) for column in frame.columns],
        "records": _records_payload(frame),
    }


def _records_payload(frame: Any) -> list[dict[str, Any]]:
    if not isinstance(frame, pd.DataFrame) or frame.empty:
        return []
    safe_frame = frame.copy()
    safe_frame.columns = [str(column) for column in safe_frame.columns]
    return [
        {str(key): _json_safe(value) for key, value in row.items()}
        for row in safe_frame.to_dict(orient="records")
    ]


def _json_safe(value: Any) -> Any:
    if isinstance(value, Mapping):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, list | tuple | set):
        return [_json_safe(item) for item in value]
    if isinstance(value, pd.DataFrame):
        return _frame_payload(value)
    if hasattr(value, "item"):
        try:
            return value.item()
        except (TypeError, ValueError):
            return str(value)
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass
    if isinstance(value, Path):
        return str(value)
    return value


def _resolve_run_root(artifacts: Mapping[str, Any]) -> Path | None:
    output_root = artifacts.get("output_root")
    if output_root:
        path = Path(str(output_root))
        if path.exists():
            return path
    manifest = artifacts.get("manifest") or artifacts.get("artifact_manifest")
    if manifest:
        path = Path(str(manifest))
        if path.exists():
            return path.parent
    return None


def _source_arcname(path: Path, run_root: Path | None) -> str:
    relative = _relative_source_path(path, run_root)
    return f"{SOURCE_ARTIFACTS_DIR}/{_posix(relative)}"


def _relative_source_path(path: Path, run_root: Path | None) -> Path:
    if run_root is not None:
        try:
            return path.resolve().relative_to(run_root.resolve())
        except ValueError:
            pass
    return Path(_safe_name(path.name))


def _posix(path: Path) -> str:
    return path.as_posix().replace("..", "_")


def _safe_name(value: str) -> str:
    safe = "".join(
        character if character.isalnum() or character in "._-" else "_"
        for character in value
    )
    return safe.strip("._") or "artifact"


def _skip_record(
    *,
    key: str,
    original_path: Path,
    reason: str,
    size_bytes: int | None = None,
) -> dict[str, Any]:
    return {
        "artifact_key": key,
        "original_path": str(original_path),
        "reason": reason,
        "size_bytes": _file_size(original_path) if size_bytes is None else size_bytes,
    }


def _file_size(path: Path) -> int:
    try:
        return int(path.stat().st_size)
    except OSError:
        return 0


def _mapping(value: Any) -> Mapping[str, Any]:
    return value if isinstance(value, Mapping) else {}


def _looks_row_level_path(path: Path) -> bool:
    path_text = path.as_posix().lower()
    row_level_tokens = (
        "/data/",
        "input_snapshot",
        "predictions",
        "scored",
        "row_level",
        "full_data",
    )
    return any(token in path_text for token in row_level_tokens)


def _looks_row_level_table(
    name: str,
    table: pd.DataFrame,
    *,
    total_rows: int | None = None,
) -> bool:
    name_lower = name.lower()
    if any(token in name_lower for token in ("prediction", "input_snapshot", "row_level")):
        return True
    columns = {str(column).lower() for column in table.columns}
    score_columns = {
        "predicted_probability",
        "predicted_class",
        "predicted_value",
        "prediction_score",
        "residual",
    }
    identifier_like = {"loan_id", "account_id", "customer_id", "entity_id", "split"}
    row_count = int(total_rows) if total_rows is not None else len(table)
    return bool(score_columns & columns) and bool(identifier_like & columns) and row_count > 500


def _evidence_area_for_path(path: Path, key: str) -> str:
    text = path.as_posix().lower()
    if key in {"config", "runner_script", "rerun_readme"} or "/config/" in text or "/code/" in text:
        return "reproducibility"
    if "/metadata/" in text:
        return "metadata"
    if "/tables/" in text:
        return "diagnostic_tables"
    if "/figures/" in text:
        return "visual_evidence"
    if "/reports/" in text:
        return "narrative_reports"
    if "/model/" in text:
        return "model_evidence"
    if path.name == "START_HERE.md":
        return "orientation"
    return "supporting_evidence"


def _llm_use_for_path(path: Path, key: str) -> str:
    file_name = path.name
    use_by_name = {
        "run_config.json": "Resolved run configuration for methodology and control documentation.",
        "generated_run.py": (
            "Generated Python workflow for reproducibility and implementation description."
        ),
        "decision_summary.md": (
            "Decision recommendation, metric scorecard, issues, and evidence index."
        ),
        "model_development_dossier.md": "Audit-oriented narrative source for methodology drafting.",
        "validation_pack.md": "Validator-facing evidence source for model validation sections.",
        "model_documentation_pack.md": "Development-facing documentation source.",
        "metrics.json": "Structured model performance metrics.",
        "statistical_tests.json": "Structured statistical-test evidence.",
        "reproducibility_manifest.json": (
            "Environment, package, hash, and input-source audit evidence."
        ),
        "artifact_manifest.json": "Authoritative run artifact index.",
        "feature_lineage_map.csv": (
            "Feature source, transformation, imputation, and documentation lineage."
        ),
    }
    if file_name in use_by_name:
        return use_by_name[file_name]
    if key == "figures_dir":
        return "Visual evidence to interpret charts and model diagnostics."
    if key == "tables_dir":
        return "Diagnostic table evidence to support quantitative claims."
    return "Supporting run evidence for model-document drafting."
